from docx import Document
from docx.enum.section import WD_SECTION
import os

from Writing_text import text, text_writing, layout
from Reading_text import text_reading


from docx.enum.style import WD_STYLE_TYPE

from classes import *

def main():
    # file name of the report
    report_file = 'report.docx'

    # create document
    report = Document()

    # name of the directory and the input files
    input_directory = 'Inputs'
    text_input_file = 'Text_input.docx'
    definitions_file = 'Terms_definitions.docx'

    # path of the input files
    text_input_path = text_reading.get_path(text_input_file, input_directory)
    definitions_path = text_reading.get_path(definitions_file, input_directory)

    # load text input files with python-docx
    text_input = Document(text_input_path)
    definitions = Document(definitions_path)

    # list of all tables in text input
    tables = [
        'Report table',
        'Study table',
        'Header table',
        'Approval table',
        'Purpose parameter table',
        'Background parameter table',
        'Scope parameter table',
        'EU Regulation definition table',
        'IEC 62366-1 definition table',
        'FDA Guidance definition table',
        'Ethics statement parameter table',
        'Device specifications parameter table',
        'Goal parameter table',
        'Participants number table',
        'Participants description table',
        'Participants parameter table',
        'Use environment parameter table',
        'Critical tasks number table',
        'Critical tasks description table',
        'Use scenarios parameter table',
        'Setup parameter table',
        'Effectiveness analysis tasks and problems table',
        'Effectiveness analysis problem number table',
        'Effectiveness analysis problem type table',
        'Effectiveness analysis video table',
        'Time on tasks table',
        'Conclusion parameter table',
    ]

    parameters = {}

    # text_reading.parameters_from_standard_tables(text_input, tables.index('Report table'), parameters)

    text_reading.get_parameters_from_tables(text_input_path, tables, parameters)

    print(parameters)


    # define all styles used in the document
    layout.define_all_styles(report)

    # section1 includes cover page and table of content
    section1 = report.sections[0]
    layout.define_page_format(section1)

    # add title
    title = text_writing.write_title(report, text.title)

    # add subtitle
    subtitle = report.add_heading(text.subtitle, 0)
    subtitle.style = report.styles['Subtitle']

    # create table of the document approval
    approval_table = report.add_table(rows=4, cols=4)

    # approval_table.style = 'Table Grid'
    layout.define_table_style(approval_table)
    for i in range(0, 4):
        for j in range(0, 4):
            approval_table.cell(i, j).text = text.approval_cells[i, j]

    # set the shading of the first row to light_grey_10 (RGB Hex: D0CECE)
    for cell in approval_table.rows[0].cells:
        layout.set_cell_shading(cell, 'D0CECE')

    # make the first row bold
    for col in approval_table.columns:
        col.cells[0].paragraphs[0].runs[0].font.bold = True

    # add the function to the person
    approval_table.cell(1, 1).add_paragraph(text.function_author)
    approval_table.cell(2, 1).add_paragraph(text.function_reviewer)
    approval_table.cell(3, 1).add_paragraph(text.function_approver)

    # make the function italic
    for i in range(1, 4):
        approval_table.rows[i].cells[1].paragraphs[1].runs[0].font.italic = True

    # add a page break
    report.add_page_break()

    # add table of content
    table_of_content = report.add_heading(text.toc_title, 1)

    # add a new section
    section2 = report.add_section(WD_SECTION.NEW_PAGE)
    layout.define_page_format(section2)

    # add a header
    header = layout.create_header(section2)
    text_writing.write_header(header.paragraphs[0], header.paragraphs[1], text.first_header, text.second_header)

    '''  create and write all the chapters  '''

    purpose = Chapter(report, text_input_path, 'Purpose', tables, parameters)
    purpose.write_chapter()

    background = Chapter(report, text_input_path, 'Background', tables, parameters)
    background.write_chapter()

    scope = Chapter(report, text_input_path, 'Scope', tables, parameters)
    scope.write_chapter()

    text_writing.write_definitions_chapter(report, text.definitions_title, text.defined_terms, text.definitions_list,
                                           text.definitions_styles_list)

    ethics = Chapter(report, text_input_path, 'Ethics statement', tables, parameters)
    ethics.write_chapter()

    device = Chapter(report, text_input_path, 'Device specifications', tables, parameters)
    device.write_chapter()

    report.add_heading(text.procedure_title, 1)

    goal = Chapter(report, text_input_path, 'Goal', tables, parameters)
    goal.write_chapter()

    # text_writing.write_chapter(report, text.goal_title, 2, text.goal_paragraphs)

    participants = Chapter(report, text_input_path, 'Participants', tables, parameters)
    participants.write_chapter()

    environment = Chapter(report, text_input_path, 'Use environment', tables, parameters)
    environment.write_chapter()

    scenarios = Chapter(report, text_input_path, 'Use scenarios', tables, parameters)
    scenarios.write_chapter()

    setup = Chapter(report, text_input_path, 'Setup', tables, parameters)
    setup.write_chapter()

    report.add_heading(text.results_title, 1)

    effectiveness_analysis = Results(report, text_input_path, 'Effectiveness analysis', tables, parameters)
    effectiveness_analysis.visualization()

    # text_writing.write_chapter(report, text.results_title, 1, text.results_paragraphs)

    conclusion = Chapter(report, text_input_path, 'Conclusion', tables, parameters)
    conclusion.write_chapter()

    # save the report
    report.save(report_file)

    # open the report with the default handler for .docx (Word)
    os.startfile(report_file)


if __name__ == '__main__':
    main()
