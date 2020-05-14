#####     TEXT ATTRIBUTION     #####

# All texts that should be writen in the report is stored in variables in this module.
# It includes: title, subtitles, paragraphs, tables entries, ...

import numpy as np
from datetime import date
from docx import Document

from Reading_text import text_reading



# name of the directory and the input files
input_directory = 'Inputs'
text_input_file = 'Text_input.docx'
definitions_file = 'Terms_definitions.docx'

# path of the input files
text_input_path = text_reading.get_path(text_input_file, input_directory)
definitions_path = text_reading.get_path(definitions_file, input_directory)

# load text input files with python-docx
text_input = Document(text_input_path)
definitions = Document(definitions_path)

'''
# attribution of all tables in the text input file
tab_report = text_input.tables[0]
tab_study = text_input.tables[1]
tab_header = text_input.tables[2]
tab_approval = text_input.tables[3]
tab_purpose_param =
tab_purpose_pic =
tab_background_param =
tab_background_pic =
tab_scope_param =
tab_scope_pic =
tab_ethics_param =
tab_ethics_pic =
tab_device_param =
tab_device_pic =
tab_goal_param =
tab_goal_pic =
'''

# report information table
tab_report = text_input.tables[0]
title = tab_report.cell(0, 1).text
subtitle = tab_report.cell(1, 1).text


# table entries of the document approval table
tab_approval = text_input.tables[4]
name_author = tab_approval.cell(0, 1).text
function_author = tab_approval.cell(1, 1).text
name_reviewer = tab_approval.cell(2, 1).text
function_reviewer = tab_approval.cell(3, 1).text
name_approver = tab_approval.cell(4, 1).text
function_approver = tab_approval.cell(5, 1).text
approval_cells = np.array((['Role', 'Name / Function', 'Date', 'Signature'],
                           ['Author', name_author, '', ''],
                           ['Reviewer', name_reviewer, '', ''],
                           ['Approver', name_approver, '', '']))

# table of content title
toc_title = 'Table of content'


# today's date
today_date = date.today()
date_str = today_date.strftime('%d.%m.%Y')

# table entries of the header in the second section
tab_header = text_input.tables[2]
header_firm = tab_header.cell(0,1).text
header_title = tab_header.cell(1,1).text
header_ID = tab_header.cell(2,1).text

# header in the second section
first_header = '{} \t {} \t {}'.format(header_firm, header_title, header_ID)
second_header = ' \t \t {}'.format(date_str)

tab_definition = text_input.tables[3]

# store all terms that we want to define in a list
defined_terms =[]
text_reading.find_definitions(tab_definition, [1, 4], defined_terms)

# store the definitions paragraphs and their styles in lists
definitions_list = []
definitions_styles_list = []
for i in range(len(defined_terms)):
    new_definition = []
    new_styles = []
    text_reading.paragraph_after_heading_with_styles(definitions.paragraphs, new_definition, new_styles, defined_terms[i], 'Heading 2')
    definitions_list.append(new_definition)
    definitions_styles_list.append(new_styles)


purpose_title = 'Purpose'
purpose_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, purpose_paragraphs, purpose_title, 'Heading 1')

background_title = 'Background'
background_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, background_paragraphs, background_title, 'Heading 1')

scope_title = 'Scope'
scope_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, scope_paragraphs, scope_title, 'Heading 1')

definitions_title = 'Terms and abbreviations'

ethics_title = 'Ethics statement'
ethics_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, ethics_paragraphs, ethics_title, 'Heading 1')

device_title = 'Device specification'
device_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, device_paragraphs, device_title, 'Heading 1')

procedure_title = 'Test procedure'
procedure_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, procedure_paragraphs, procedure_title, 'Heading 1')

goal_title = 'Goal'
goal_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, goal_paragraphs, goal_title, 'Heading 2')

participants_title = 'Participants'
participants_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, participants_paragraphs, participants_title, 'Heading 2')

environment_title = 'Use environment'
environment_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, environment_paragraphs, environment_title, 'Heading 2')

scenarios_title = 'Use scenarios / Use cases'
scenarios_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, scenarios_paragraphs, scenarios_title, 'Heading 2')

setup_title = 'Set up'
setup_paragraphs = []
text_reading.paragraph_after_heading_different(text_input.paragraphs, setup_paragraphs, setup_title, 'Heading 2', 'Heading 1')

results_title = 'Results'
results_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, results_paragraphs, results_title, 'Heading 1')

conclusion_title = 'Conclusion'
conclusion_paragraphs = []
text_reading.paragraph_after_heading(text_input.paragraphs, conclusion_paragraphs, conclusion_title, 'Heading 1')


# read dropdown lists from text input file
dd_lists_values = []
text_reading.read_dropdown_lists(text_input_file, dd_lists_values)



'''
paragraphs = list(text_input.paragraphs)

paragraphs_position = []

list_of_title = []

for i in range(len(paragraphs)):
    if paragraphs[i].style.name == 'Heading 1':
        if paragraphs[i].text == 'Purpose':
            purpose_title = paragraphs[i].text
            list_of_title.append(paragraphs[i])
            purpose_pos = i
            paragraphs_position.append(i)
        elif paragraphs[i].text == 'Background':
            background_title = paragraphs[i].text
            list_of_title.append(paragraphs[i])
            background_pos = i
            paragraphs_position.append(i)
        elif paragraphs[i].text == 'Scope':
            scope_title = paragraphs[i].text
            list_of_title.append(paragraphs[i])
            scope_pos = i
            paragraphs_position.append(i)
        elif paragraphs[i].text == 'Ethics statement':
            ethics_title = paragraphs[i].text
            list_of_title.append(paragraphs[i])
            ethics_pos = i
            paragraphs_position.append(i)
        elif paragraphs[i].text == 'Device specification':
            device_title = paragraphs[i].text
            list_of_title.append(paragraphs[i])
            device_pos = i
            paragraphs_position.append(i)
        elif paragraphs[i].text == 'Test procedure':
            procedure_title = paragraphs[i].text
            list_of_title.append(paragraphs[i])
            procedure_pos = i
            paragraphs_position.append(i)
        elif paragraphs[i].text == 'Results':
            results_title = paragraphs[i].text
            list_of_title.append(paragraphs[i])
            results_pos = i
            paragraphs_position.append(i)
        elif paragraphs[i].text == 'Conclusion':
            conclusion_title = paragraphs[i].text
            list_of_title.append(paragraphs[i])
            conclusion_pos = i
            paragraphs_position.append(i)

list_of_paragraphs = []

for i in range(len(paragraphs_position) - 1):
    list_of_paragraphs.append(paragraphs[paragraphs_position[i] + 1: paragraphs_position[i + 1]])
'''