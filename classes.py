from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from zipfile import ZipFile
from bs4 import BeautifulSoup

from Writing_text import layout
from Reading_text import text_reading

class File:
    ''' Class that defines a file '''

    def __init__(self, file_name, directory_name):
        self.name = file_name
        self.directory = directory_name

    # return the path of a file given its name and the name of its directory
    @ property
    def path(self):
        text_input_directory = os.path.dirname(self.name)
        path = os.path.join(text_input_directory, self.directory, self.name)

        return path


class Chapter:
    '''A class that defines a chapter ...'''

    TEXT_INPUT_FILE = 'Text_input.docx'
    REPORT_FILE = 'Report.docx'
    DEFINITIONS_FILE = 'Terms_definitions.docx'
    INPUTS_DIRECTORY = 'Inputs'

    def __init__(self, report_document, text_input_path, title, list_of_tables, parameters_dictionary):
        self.report = report_document
        self.text_input_path = text_input_path
        self.text_input = Document(text_input_path)
        self.title = title
        self.list_of_tables = list_of_tables
        self.parameters_dictionary = parameters_dictionary
        # self.heading_level = heading_level
        # self.parameter_table_index = parameter_table_index
        # self.picture_table = picture_table

    '''
    @ property
    def text_input(self):
        text_input = File(self.TEXT_INPUT_FILE, self.INPUTS_DIRECTORY)
        return Document(text_input.path)
    '''

    #  find a heading given his title and return the corresponding paragraph index
    def heading_name_index(self):
        for i in range(len(self.text_input.paragraphs)):
            if self.text_input.paragraphs[i].style.name == 'Heading 1':
                if self.text_input.paragraphs[i].text == self.title:
                    return i
            elif self.text_input.paragraphs[i].style.name == 'Heading 2':
                if self.text_input.paragraphs[i].text == self.title:
                    return i
            elif self.text_input.paragraphs[i].style.name == 'Heading 3':
                if self.text_input.paragraphs[i].text == self.title:
                    return i
            elif self.text_input.paragraphs[i].style.name == 'Heading 4':
                if self.text_input.paragraphs[i].text == self.title:
                    return i

            '''
            if self.text_input.paragraphs[i].style.name == 'Heading {}'.format(1 or 2 or 3 or 4):  # look for paragraphs with corresponding style
                if self.text_input.paragraphs[i].text == self.title:  # look for paragraphs with corresponding title
                    return i  # return the index of the paragraphs
            '''

    # return the index of the next heading corresponding to a style given the index of the previous heading
    def next_heading_index(self, previous_index):
        for i in range(previous_index + 1, len(self.text_input.paragraphs)):
            if self.text_input.paragraphs[i].style.name == 'Heading 1':
                return i
            elif self.text_input.paragraphs[i].style.name == 'Heading 2':
                return i
            elif self.text_input.paragraphs[i].style.name == 'Heading 3':
                return i
            elif self.text_input.paragraphs[i].style.name == 'Heading 4':
                return i

            '''
            if self.text_input.paragraphs[i].style.name == 'Heading {}'.format(1 or 2 or 3 or 4):  # look for paragraphs with corresponding style
                return i  # return the index of the paragraph
            '''

    # store all paragraphs and their corresponding style between a given heading and the next one
    def paragraph_after_heading_with_styles(self, list_of_paragraphs, list_of_styles):
        heading_index = self.heading_name_index()  # index of the given heading
        next_index = self.next_heading_index(heading_index)  # index of the next heading
        for i in range(heading_index + 1, next_index):  # loop over all paragraphs between the given heading and the next one
            list_of_paragraphs.append(self.text_input.paragraphs[i])  # store all paragraphs in a given list
            list_of_styles.append(self.text_input.paragraphs[i].style.name)  # store all styles in a given list

    # return the paragraphs following the given heading
    @ property
    def paragraphs(self):
        list_of_paragraphs = []
        heading_index = self.heading_name_index()  # index of the given heading
        next_index = self.next_heading_index(heading_index)  # index of the next heading
        for i in range(heading_index + 1, next_index):  # loop over all paragraphs between the given heading and the next one
            list_of_paragraphs.append(self.text_input.paragraphs[i])  # store all paragraphs in a given list

        return list_of_paragraphs

    @ property
    def parameters(self):
        """
        Read the dropdown lists of the parameter table and return their value in a list.
        """

        return text_reading.get_dropdown_list_of_table(self.text_input_path,
                                                       self.list_of_tables.index('{} parameter table'.format(self.title))
                                                       )

    def write_chapter(self):
        """
        Write the heading and the paragraphs of a chapter, including the parameters.
        """

        # read heading style and write heading
        heading_style = self.text_input.paragraphs[self.heading_name_index()].style.name
        self.report.add_paragraph(self.title, heading_style)

        parameters_values = ['', '', '']

        # stores values of corresponding parameter keys in a list as lower case string
        for i in range(len(self.parameters)):
            if self.parameters[i] != '-':
                parameters_values[i] = self.parameters_dictionary[self.parameters[i]].lower()

        # write paragraphs including values of parameters
        for i in range(len(self.paragraphs)):
            paragraph = self.report.add_paragraph(
                self.paragraphs[i].text.format(parameters_values[0], parameters_values[1], parameters_values[2])
            )
            paragraph.style.name = 'Normal'


class Results:
    def __init__(self, report_document, text_input_path, title, list_of_tables, parameters_dictionary):
        self.report = report_document
        self.text_input_path = text_input_path
        self.text_input = Document(text_input_path)
        self.title = title
        self.list_of_tables = list_of_tables
        self.parameters_dictionary = parameters_dictionary

    def visualization(self):
        # index of the input tables
        task_table_index = self.list_of_tables.index('Effectiveness analysis tasks and problems table')
        problem_table_index = self.list_of_tables.index('Effectiveness analysis problem type table')
        video_table_index = self.list_of_tables.index('Effectiveness analysis video table')

        # input tables
        task_table = self.text_input.tables[task_table_index]
        problem_table = self.text_input.tables[problem_table_index]
        video_table = self.text_input.tables[video_table_index]

        # color
        light_grey_10 = 'D0CECE'
        green = '00B050'
        red = 'FF0000'
        orange = 'FFC000'
        yellow = 'FFFF00'

        # create a table for the results visualization
        result_table_rows = self.parameters_dictionary['Number of critical tasks'] + 1
        result_table_cols = self.parameters_dictionary['Number of participants'] + 1
        result_table = self.report.add_table(result_table_rows, result_table_cols)
        layout.define_table_style(result_table)

        # write the information of the input table in the result table
        for i in range(result_table_rows):
            for j in range(result_table_cols):
                cell = result_table.cell(i, j)

                # skip the first row and first column
                if i != 0 and j != 0:
                    cell.text = task_table.cell(i, j).text
                    cell.paragraphs[0].runs[0].font.bold = True

                # first row
                elif i == 0 and j != 0:
                    cell.text = task_table.cell(i, j).text
                    layout.set_cell_shading(cell, light_grey_10)     # color the cell in light_grey_10
                    cell.paragraphs[0].runs[0].font.bold = True

                # first column
                elif i != 0 and j == 0:
                    cell.text = self.parameters_dictionary['Critical task {} name'.format(i)]
                    layout.set_cell_shading(cell, light_grey_10)     # color the cell in light_grey_10
                    cell.paragraphs[0].runs[0].font.bold = True

                # bolds all text
                if cell.text:
                    cell.paragraphs[0].runs[0].font.bold = True

        # color the cell according to the type of problem
        for i in range(1, result_table_rows):
            for j in range(1, result_table_cols):
                cell = result_table.cell(i, j)

                if cell.text:     # check if the text string is not empty
                    index = int(cell.text)

                    if self.parameters_dictionary['Problem {} type'.format(index)] == 'Important problem':
                        layout.set_cell_shading(cell, orange)

                    if self.parameters_dictionary['Problem {} type'.format(index)] == 'Marginal problem':
                        layout.set_cell_shading(cell, yellow)

                    if self.parameters_dictionary['Problem {} type'.format(index)] == 'Critical problem':
                        layout.set_cell_shading(cell, red)
                else:
                    layout.set_cell_shading(cell, green)
