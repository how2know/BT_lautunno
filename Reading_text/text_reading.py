#####     TEXT READING MODULE     #####

# Functions that read the text from the text input file are implemented in this module.

from docx import Document
import os
from zipfile import ZipFile
from bs4 import BeautifulSoup


# return the path of a file given its name and the name of its directory
def get_path(file_name, directory_name):
    text_input_directory = os.path.dirname(file_name)
    path = os.path.join(text_input_directory, directory_name, file_name)

    return path


#  find a heading with his title and style and return the corresponding paragraph index
def heading_name_index(paragraphs, title, style):
    for i in range(len(paragraphs)):              # loop over all paragraphs
        if paragraphs[i].style.name == style:     # look for paragraphs with corresponding style
            if paragraphs[i].text == title:       # look for paragraphs with corresponding title
                return i                          # return the index of the paragraphs


# return the index of the next heading corresponding to a style given the index of the previous heading
def next_heading_index(paragraphs, style, previous_index):
    for i in range(previous_index + 1, len(paragraphs)):     # loop over paragraphs coming after the given paragraph index
        if paragraphs[i].style.name == style:                # look for paragraphs with corresponding style
            return i                                         # return the index of the paragraph


# return the index of the next heading corresponding to a style given the index of the previous heading
def next_different_heading_index(paragraphs, style, previous_index):
    for i in range(previous_index + 1, len(paragraphs)):     # loop over paragraphs coming after the given paragraph index
        if paragraphs[i].style.name == style:                # look for paragraphs with corresponding style
            return i                                         # return the index of the paragraph


# store all paragraphs and their corresponding style between a given heading and the next one
def paragraph_after_heading_with_styles(paragraphs, list_of_paragraphs, list_of_styles, heading_title, heading_style):
    heading_index = heading_name_index(paragraphs, heading_title, heading_style)     # index of the given heading
    next_index = next_heading_index(paragraphs, heading_style, heading_index)        # index of the next heading
    for i in range(heading_index + 1, next_index):                                   # loop over all paragraphs between the given heading and the next one
        list_of_paragraphs.append(paragraphs[i])                                     # store all paragraphs in a given list
        list_of_styles.append(paragraphs[i].style.name)                              # store all styles in a given list


# store all paragraphs between a given heading and the next one
def paragraph_after_heading(paragraphs, list_of_paragraphs, heading_title, heading_style):
    heading_index = heading_name_index(paragraphs, heading_title, heading_style)     # index of the given heading
    next_index = next_heading_index(paragraphs, heading_style, heading_index)        # index of the next heading
    for i in range(heading_index + 1, next_index):                                   # loop over all paragraphs between the given heading and the next one
        list_of_paragraphs.append(paragraphs[i])                                     # store all paragraphs in a given list


# store all paragraphs between a given heading and the next one
def paragraph_after_heading_different(paragraphs, list_of_paragraphs, heading_title, heading_style1, heading_style2):
    heading_index = heading_name_index(paragraphs, heading_title, heading_style1)     # index of the given heading
    next_index = next_heading_index(paragraphs, heading_style2, heading_index)        # index of the next heading
    for i in range(heading_index + 1, next_index):                                   # loop over all paragraphs between the given heading and the next one
        list_of_paragraphs.append(paragraphs[i])                                     # store all paragraphs in a given list


# find all terms that you want to define and store them
def find_definitions(table, columns_indexes_list, definitions_list):
    for j in columns_indexes_list:                                     # loop over the columns that contains "Yes" or "No"
        for i in range(len(table.columns[j].cells)):                   # loop over all cells of the columns
            if table.cell(i, j).text == 'Yes':                         # find all cells that contains "Yes"
                definitions_list.append(table.cell(i, j - 1).text)     # store the terms that correspond to a "Yes"


# read dropdown lists and store their value in a list
def read_dropdown_lists(file_name, list_of_value):

    # open docx file as a zip file and store its relevant xml data
    zip_file = ZipFile(file_name)
    xml_data = zip_file.read('word/document.xml')
    zip_file.close()

    # parse the xml data with BeautifulSoup
    soup = BeautifulSoup(xml_data, 'xml')

    # look for all values of dropdown lists in the data and store them
    dd_lists_content = soup.find_all('sdtContent')
    for i in dd_lists_content:
        list_of_value.append(i.find('t').string)


# return a list of the value of all dropdown lists in a table
def get_dropdown_list_of_table(text_input_path, table_index):
    list_of_value = []

    # open docx file as a zip file and store its relevant xml data
    zip_file = ZipFile(text_input_path)
    xml_data = zip_file.read('word/document.xml')
    zip_file.close()

    # parse the xml data with BeautifulSoup
    soup = BeautifulSoup(xml_data, 'xml')

    # look for all values of dropdown lists in the data and store them
    tables = soup.find_all('tbl')
    dd_lists_content = tables[table_index].find_all('sdtContent')
    for i in dd_lists_content:
        list_of_value.append(i.find('t').string)

    return list_of_value


#
def get_parameters_from_tables(text_input_path, list_of_table, parameters_dictionary):
    '''
    Function that reads all parameters stored in the tables of the text input document and stores them in a dictionary.

    Standard tables that define the parameters have two columns, which contain the key and the value of parameters.
    The parameters of these tables are read in the same way for all tables, the only difference is made when the value
    should be an integer or a string.

    Special tables differs from standard tables, e.g. they have more than two columns, contain dropdown lists, etc...
    The parameters of these tables are read are read in a different way for all tables.

    :param text_input_path:
    :param list_of_table:
    :param parameters_dictionary:
    :return:
    '''

    text_input_document = Document(text_input_path)

    # tables with two columns having parameters in each row
    # those tables are handled the same way
    standard_parameters_table = [
        'Report table',
        'Study table',
        'Header table',
        'Approval table',
        'Participants number table',
        'Critical tasks number table',
        'Effectiveness analysis problem number table',
    ]

    # tables that have special features (e.g. more than two columns, dropdown lists, ...)
    # those tables are handled separately
    special_parameters_table = [
        'Critical tasks description table',
        'Effectiveness analysis problem type table'
    ]

    # get the parameters from the standard table
    for table_name in standard_parameters_table:
        table_index = list_of_table.index(table_name)
        table = text_input_document.tables[table_index]

        for row in table.rows:
            key = row.cells[0].text

            if key.startswith('Number of'):
                parameters_dictionary[key] = int(row.cells[1].text)
            else:
                parameters_dictionary[key] = row.cells[1].text

    # get the parameters from the special table
    for table_name in special_parameters_table:

        # get the parameters from the critical tasks description table
        if table_name.startswith('Critical tasks description'):
            table_index = list_of_table.index(table_name)
            table = text_input_document.tables[table_index]

            for i in range(1, parameters_dictionary['Number of critical tasks'] + 1):
                type_key = table.cell(i, 0).text + ' name'
                description_key = table.cell(i, 0).text + ' description'

                parameters_dictionary[type_key] = table.cell(i, 1).text
                parameters_dictionary[description_key] = table.cell(i, 2).text

        # get the parameters from the effectiveness analysis problem type table
        if table_name.startswith('Effectiveness analysis problem'):
            table_index = list_of_table.index(table_name)
            table = text_input_document.tables[table_index]

            '''Cells are not recognized as cells by word if they contain a dropdown list. That is why I had to create 
            a work around to get the values here.'''

            list_of_text = []
            stop = False

            # stores the text of every relevant cell in a list
            while not stop:
                for i in range(1, 5 + 1):
                    for cell in table.rows[i].cells:
                        text = cell.text

                        if not text:
                            list_of_text.pop()
                            stop = True
                        else:
                            list_of_text.append(text)

            # delete the last item to ensure that there is no key without a corresponding value
            if len(list_of_text) % 2 != 0:
                list_of_text.pop()

            dropdown_lists_values = get_dropdown_list_of_table(text_input_path, table_index)

            value = 0

            # stores the keys and corresponding values in the dictionary
            for i in range(len(list_of_text)):
                if i % 2 == 0:
                    problem_number = list_of_text[i]
                    description = list_of_text[i + 1]

                    type_key = problem_number + ' type'
                    description_key = problem_number + ' description'

                    parameters_dictionary[type_key] = dropdown_lists_values[value]
                    parameters_dictionary[description_key] = description

                    value += 1