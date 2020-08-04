from docx import Document
from docx.enum.section import WD_SECTION
import os
import win32com.client
import inspect
import time

from docx_package.layout import Layout
from docx_package.chapter import Chapter
from docx_package.effectiveness_analysis import EffectivenessAnalysis
from docx_package.definitions import Definitions
from docx_package.header_footer import Header, Footer
from docx_package.table_of_content import TableOfContent
from docx_package.time_on_tasks import TimeOnTasks
from docx_package.cover_page import CoverPage
from docx_package.dwell_times_revisits import DwellTimesAndRevisits
from docx_package.average_fixation import AverageFixation
from docx_package.transitions import Transitions
from docx_package.parameters import Parameters
from docx_package.picture import Picture
from docx_package.document_history import DocumentHistory
from docx_package.participants_characteristics import ParticipantsCharacteristics
from docx_package.dropdown_lists import DropDownLists
from docx_package.use_scenarios import UseScenarios

from eye_tracking_package.cGOM_data import cGOM
from eye_tracking_package.tobii_data import TobiiData


def update(report_file):
    """
    Update all fields of  the report, i.e. the table of content, the figure captions and the list of figures.

    Args:
        report_file: File name of the report.
    """

    # get the absolut path of the report
    script_dir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
    report_file_path = os.path.join(script_dir, report_file)

    # open the report through Word, update all fields, save and quit Word
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(report_file_path)
    word.ActiveDocument.Fields.Update()
    doc.Close(SaveChanges=True)
    word.Quit()


def main():

    main_start = time.time()

    # file name of the report
    report_file = 'Report.docx'

    # create the report document
    report = Document()

    # path of the input files
    text_input_path = 'Inputs/Text_input_test2.docx'
    definitions_path = 'Inputs/Terms_definitions.docx'

    # load text input files with python-docx
    text_input = Document(text_input_path)
    definitions = Document(definitions_path)

    # path to the pictures that must be added to the report
    picture_paths = Picture.get_picture_paths()

    # soup of the text input form document
    text_input_soup = DropDownLists.get_soup(text_input_path)

    # list of all tables in the order they appear in the text input document,
    # this is used to get the index of a table in the document
    tables = [
        'Study table',
        'Title table',
        'Approval table',
        'Cover page caption table',
        'Header table',
        'Purpose text table',
        'Purpose parameter table',
        'Purpose caption table',
        'Background text table',
        'Background parameter table',
        'Background caption table',
        'Scope text table',
        'Scope parameter table',
        'Scope caption table',
        'EU Regulation 2017/745 definitions table',
        'IEC 62366-1 definitions table',
        'FDA Guidance definitions table',
        'Ethics statement text table',
        'Ethics statement parameter table',
        'Ethics statement caption table',
        'Device specifications text table',
        'Device specifications parameter table',
        'Device specifications caption table',
        'Goal text table',
        'Goal parameter table',
        'Goal caption table',
        'Participants text table',
        'Participants parameter table',
        'Participants caption table',
        'Use environment text table',
        'Use environment parameter table',
        'Use environment caption table',
        'Use scenarios text table',
        'Use scenarios parameter table',
        'Use scenarios caption table',
        'Setup text table',
        'Setup parameter table',
        'Setup caption table',
        'Critical tasks description table',
        'Effectiveness analysis decision table',
        'Effectiveness analysis tasks and problems table',
        'Effectiveness analysis problem type table',
        'Effectiveness analysis text table',
        'Effectiveness analysis parameter table',
        'Effectiveness analysis caption table',
        'Time on tasks decision table',
        'Time on tasks plot type table',
        'Time on tasks table',
        'Time on tasks text table',
        'Time on tasks parameter table',
        'Time on tasks caption table',
        'Dwell times and revisits decision table',
        'Dwell times and revisits text table',
        'Dwell times and revisits parameter table',
        'Dwell times and revisits caption table',
        'Average fixation decision table',
        'Average fixation plot type table',
        'Average fixation text table',
        'Average fixation parameter table',
        'Average fixation caption table',
        'Transitions decision table',
        'Transitions text table',
        'Transitions parameter table',
        'Transitions caption table',
        'Conclusion text table',
        'Conclusion parameter table',
        'Conclusion caption table',
        'Participants characteristics table'
    ]

    # parameters needed to write the report
    parameters = Parameters.get_all(text_input, text_input_soup, tables)

    # list of data frames that contain the cGOM data
    cGOM_dataframes = cGOM.make_dataframes_list()

    # data frame that contain the Tobii data
    tobii_data = TobiiData.make_main_dataframe(parameters)

    # define all styles used in the document
    Layout.define_all_styles(report)

    ######   COVER PAGE   ######

    section1 = report.sections[0]
    Layout.define_page_format(section1)

    cover_page_start = time.time()
    cover_page = CoverPage(report, text_input, tables, picture_paths, parameters)
    cover_page.create()

    '''
    cover_page_end = time.time()
    print('Cover page added in %.2f seconds' % (cover_page_end - cover_page_start))
    '''

    ######   TABLE OF CONTENT   ######

    report.add_page_break()

    TableOfContent.write(report)

    ######   CHAPTERS   ######

    section2 = report.add_section(WD_SECTION.NEW_PAGE)
    Layout.define_page_format(section2)

    footer = Footer(section2)
    footer.write()

    header = Header(section2, parameters)
    header.write()

    purpose = Chapter(report, text_input, text_input_soup, 'Purpose', tables, picture_paths, parameters)
    purpose.write_chapter()

    background = Chapter(report, text_input, text_input_soup, 'Background', tables, picture_paths, parameters)
    background.write_chapter()

    scope = Chapter(report, text_input, text_input_soup, 'Scope', tables, picture_paths, parameters)
    scope.write_chapter()

    Definitions.write_all_definitions(report, text_input, text_input_soup, definitions, tables)

    ethics = Chapter(report, text_input, text_input_soup, 'Ethics statement', tables, picture_paths, parameters)
    ethics.write_chapter()

    device = Chapter(report, text_input, text_input_soup, 'Device specifications', tables, picture_paths, parameters)
    device.write_chapter()

    report.add_paragraph('Test procedure', 'Heading 1')

    goal = Chapter(report, text_input, text_input_soup, 'Goal', tables, picture_paths, parameters)
    goal.write_chapter()

    participants = Chapter(report, text_input, text_input_soup, 'Participants', tables, picture_paths, parameters)
    participants.write_chapter()

    environment = Chapter(report, text_input, text_input_soup, 'Use environment', tables, picture_paths, parameters)
    environment.write_chapter()

    scenarios = UseScenarios(report, text_input, text_input_soup, 'Use scenarios', tables, picture_paths, parameters)
    scenarios.write_chapter()

    setup = Chapter(report, text_input, text_input_soup, 'Setup', tables, picture_paths, parameters)
    setup.write_chapter()

    report.add_paragraph('Results', 'Heading 1')

    start1 = time.time()
    effectiveness_analysis = EffectivenessAnalysis(report, text_input, text_input_soup, tables, picture_paths, parameters)
    effectiveness_analysis.write_chapter()
    end1 = time.time()
    print('Effectiveness analysis: ', end1-start1)

    start2 = time.time()
    time_on_tasks = TimeOnTasks(report, text_input, text_input_soup, tables, picture_paths, parameters, tobii_data)
    time_on_tasks.write_chapter()
    end2 = time.time()
    print('Time on tasks: ', end2-start2)

    start3 = time.time()
    dwell_times_and_revisits = DwellTimesAndRevisits(report, text_input, text_input_soup, tables, picture_paths, parameters, cGOM_dataframes)
    dwell_times_and_revisits.write_chapter()
    end3 = time.time()
    print('Dwell times: ', end3-start3)
    
    start4 = time.time()
    average_fixation = AverageFixation(report, text_input, text_input_soup, tables, picture_paths, parameters, cGOM_dataframes)
    average_fixation.write_chapter()
    end4 = time.time()
    print('Average fixation: ', end4-start4)

    start5 = time.time()
    transitions = Transitions(report, text_input, text_input_soup, tables, picture_paths, parameters, cGOM_dataframes)
    transitions.write_chapter()
    end5 = time.time()
    print('Transitions: ', end5-start5)

    conclusion = Chapter(report, text_input, text_input_soup, 'Conclusion', tables, picture_paths, parameters)
    conclusion.write_chapter()

    DocumentHistory.write(report)

    ######   APPENDIX   ######

    report.add_page_break()

    report.add_paragraph('Appendix', 'Heading 1')

    Definitions.write_references(report, text_input, text_input_soup, definitions, tables)

    report.add_page_break()

    Picture.add_figures_list(report)

    report.add_page_break()

    ParticipantsCharacteristics.write(report, text_input, tables, parameters)

    # save the report
    report.save(report_file)

    # error message for the image files that were not added to the report
    '''Picture.error_message(picture_paths)'''

    start6 = time.time()
    # update the table of content
    update(report_file)
    end6 = time.time()
    print('Update: ', end6 - start6)

    # open the report with the default application for .docx (Word)
    os.startfile(report_file)

    end = time.time()
    print(end - main_start)


if __name__ == '__main__':
    main()
