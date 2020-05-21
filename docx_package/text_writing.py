# from Writing_text import layout
from docx_package import layout, text_reading, text
import docx.oxml.ns as ns
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement


# write the title and return it as a paragraph
def write_title(document, text):
    title = document.add_paragraph(text, 'Title')
    layout.insert_horizontal_border(title)

    return title


# write an header of two lines
def write_header(first_line, second_line, first_line_content, second_line_content):
    first_line.text = first_line_content
    second_line.text = second_line_content





'''
def write_chapter(document, heading_title, heading_level, paragraphs):
    document.add_heading(heading_title, heading_level)
    for i in range(len(paragraphs)):
        document.add_paragraph(paragraphs[i].text)


def write_definitions_chapter(document, heading_title, defined_terms_list, definitions_list, definitions_styles_list):
    document.add_heading(heading_title, 1)
    for i in range(len(defined_terms_list)):
        document.add_heading(defined_terms_list[i], 2)
        for j in range(len(definitions_list[i])):
            if definitions_styles_list[i][j] != 'Heading 1':
                document.add_paragraph(definitions_list[i][j].text, definitions_styles_list[i][j])
'''