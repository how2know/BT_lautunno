from docx.document import Document
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.shared import Pt
from bs4 import BeautifulSoup
from typing import List, Dict, Union

from docx_package.layout import Layout
from docx_package.results import ResultsChapter
from docx_package.dropdown_lists import DropDownLists


class EffectivenessAnalysis:
    """
    Class that represents the 'Effectiveness analysis' chapter and the visualization of its results.
    """

    # name of tables as they appear in the tables list
    TASK_TABLE = 'Effectiveness analysis tasks and problems table'
    PROBLEM_TABLE = 'Effectiveness analysis problem type table'
    DECISION_TABLE = 'Effectiveness analysis decision table'

    # parameter keys
    TASKS_NUMBER_KEY = 'Number of critical tasks'
    PARTICIPANTS_NUMBER_KEY = 'Number of participants'
    PROBLEMS_NUMBER_KEY = 'Number of problems'

    # information about the headings of this chapter
    TITLE = 'Effectiveness analysis'
    TITLE_STYLE = 'Heading 2'
    DESCRIPTION_TITLE = 'Problems description'
    DESCRIPTION_STYLE = 'Heading 3'
    DISCUSSION_TITLE = 'Discussion'
    DISCUSSION_STYLE = 'Heading 3'

    # hexadecimals of colors for cell shading
    LIGHT_GREY_10 = 'D0CECE'
    GREEN = '00B050'
    RED = 'FF0000'
    ORANGE = 'FFC000'
    YELLOW = 'FFFF00'

    # list of the colors table columns width and table rows height
    COLORS_TABLE_WIDTHS = [2, 0.5, 3.8, 1.6, 1.7, 0.5, 3.8, 2]
    COLORS_TABLE_HEIGHTS = [0.5, 0.2, 0.5]

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 text_input_soup: BeautifulSoup,
                 list_of_tables: List[str],
                 picture_paths_list: List[str],
                 parameters_dictionary: Dict[str, Union[str, int]]
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            list_of_tables: List of all table names.
            picture_paths_list: List of the path of all input pictures.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value)
        """

        self.report = report_document
        self.text_input = text_input_document
        self.text_input_soup = text_input_soup
        self.tables = list_of_tables
        self.picture_paths = picture_paths_list
        self.parameters = parameters_dictionary

    # TODO: write this in __init__
    @ property
    def task_table(self) -> Table:
        """
        Returns:
            Table of the input .docx file where the information about the effectiveness analysis are written.
        """

        task_table_index = self.tables.index(self.TASK_TABLE)
        return self.text_input.tables[task_table_index]

    @ property
    def problem_table(self) -> Table:
        """
        Returns:
            Table of the input .docx file where the description of the problems are written.
        """

        problem_table_index = self.tables.index(self.PROBLEM_TABLE)
        return self.text_input.tables[problem_table_index]

    @ property
    def tasks_number(self) -> int:
        """
        Returns:
            Biggest number of critical tasks between the one given in the 'Number of critical tasks' input
            and the one that corresponds to the effectiveness analysis input table.
        """

        tasks_number = 0

        # get the index of the last row that is filled which corresponds to the number of critical tasks
        for i in range(1, len(self.task_table.rows)):
            row_filled = False
            for cell in self.task_table.rows[i].cells[1:]:
                if cell.text:
                    row_filled = True
            if row_filled:
                tasks_number = i

        # choose the biggest number of critical tasks
        if tasks_number > self.parameters[self.TASKS_NUMBER_KEY]:
            return tasks_number
        else:
            return self.parameters[self.TASKS_NUMBER_KEY]

    @ property
    def participants_number(self) -> int:
        """
        Returns:
            Biggest number of participants between the one given in the 'Number of participants' input
            and the one that corresponds to the effectiveness analysis input table.
        """

        participants_number = 0

        # get the index of the last column that is filled which corresponds to the number of participants
        for j in range(1, len(self.task_table.columns)):
            column_completed = False
            for cell in self.task_table.columns[j].cells[1:]:
                if cell.text:
                    column_completed = True
            if column_completed:
                participants_number = j

        # choose the biggest number of participant
        if participants_number > self.parameters[self.PARTICIPANTS_NUMBER_KEY]:
            return participants_number
        else:
            return self.parameters[self.PARTICIPANTS_NUMBER_KEY]

    def make_result_table(self):
        """
        Create a table for the visualization of the effectiveness analysis.
        """

        # create a table for the results visualization
        rows_number = self.tasks_number + 1
        cols_number = self.participants_number + 1
        result_table = self.report.add_table(rows_number, cols_number)
        result_table.style = 'Table Grid'
        result_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        result_table.autofit = False

        # write the information of the input table in the result table
        for i in range(rows_number):
            for j in range(cols_number):
                cell = result_table.cell(i, j)

                # skip the first row and first column
                if i != 0 and j != 0:
                    cell.text = self.task_table.cell(i, j).text
                    cell.paragraphs[0].runs[0].font.bold = True

                # first row
                elif i == 0 and j != 0:
                    cell.text = self.task_table.cell(i, j).text
                    Layout.set_cell_shading(cell, self.LIGHT_GREY_10)     # color the cell in light_grey_10
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    cell.paragraphs[0].runs[0].font.bold = True

                # first column
                elif i != 0 and j == 0:
                    try:
                        cell.text = self.parameters['Critical task {} name'.format(i)]
                        cell.paragraphs[0].runs[0].font.bold = True

                    # case where no critical task name were given
                    except KeyError:
                        pass
                    Layout.set_cell_shading(cell, self.LIGHT_GREY_10)     # color the cell in light_grey_10

                # bolds all text
                if cell.text:
                    cell.paragraphs[0].runs[0].font.bold = True

        # color the cell according to the type of problem
        for i in range(1, rows_number):
            for j in range(1, cols_number):
                cell = result_table.cell(i, j)

                if cell.text:     # check if the text string is not empty
                    problem_index = int(cell.text)
                    try:
                        problem_type = self.parameters['Problem {} type'.format(problem_index)]

                    # case where no problem type were given
                    except KeyError:
                        problem_type = ''

                    if problem_type == 'Important problem':
                        Layout.set_cell_shading(cell, self.ORANGE)

                    if problem_type == 'Marginal problem':
                        Layout.set_cell_shading(cell, self.YELLOW)

                    if problem_type == 'Critical problem':
                        Layout.set_cell_shading(cell, self.RED)
                else:
                    Layout.set_cell_shading(cell, self.GREEN)

        # color the top left cell borders in white
        Layout.set_cell_border(result_table.cell(0, 0),
                               top={"color": "#FFFFFF"},
                               start={"color": "#FFFFFF"}
                               )

        # set the vertical and horizontal alignment of all cells
        for row in result_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for i in range(len(result_table.rows)):
            for j in range(1, len(result_table.columns)):
                result_table.cell(i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # set the width of the columns
        Layout.set_column_width(result_table.columns[0], 2.4)
        for j in range(1, cols_number):
            width = (15.9 - 2.4) / (cols_number - 1)
            Layout.set_column_width(result_table.columns[j], width)

        # set the height of the rows
        Layout.set_row_height(result_table.rows[0], 0.5)
        for i in range(1, rows_number):
            Layout.set_row_height(result_table.rows[i], 1.1, rule=WD_ROW_HEIGHT_RULE.AT_LEAST)

    @ staticmethod
    def add_color_description(table: Table, cell_row: int, cell_column: int, color: str, description: str):
        """
        Color the cell of a table and add a description in the following cell.

        This function is used to create the table that describes the color of the effectiveness analysis table.

        Args:
            table: Table that will contain the colors and their description.
            cell_row: Row in which the color cell is located.
            cell_column: Column in which the color cell is located.
            color: Color of the cell.
            description: Description / meaning of the color.
        """

        color_cell = table.cell(cell_row, cell_column)
        description_cell = table.cell(cell_row, cell_column + 1)

        # color the cell and set its borders
        Layout.set_cell_shading(color_cell, color)
        Layout.set_cell_border(color_cell,
                               top={"sz": 4, "val": "single", "color": "#000000"},
                               bottom={"sz": 4, "val": "single", "color": "#000000"},
                               start={"sz": 4, "val": "single", "color": "#000000"},
                               end={"sz": 4, "val": "single", "color": "#000000"},
                               )

        # write description next to the color cell and set the font size
        description_cell.text = description
        description_cell.paragraphs[0].runs[0].font.size = Pt(9)
        description_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def make_colors_table(self):
        """
        Create a table that describes the color of the effectiveness analysis table.
        """

        # add a line for space between the two tables
        separation_line = self.report.add_paragraph(' ')
        separation_line.runs[0].font.size = Pt(5)

        # create table
        colors_table = self.report.add_table(3, 8)
        colors_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        colors_table.autofit = False

        # set the height of all rows
        for idx, row in enumerate(colors_table.rows):
            Layout.set_row_height(row, self.COLORS_TABLE_HEIGHTS[idx])

        # set the width of all columns
        for idx, column in enumerate(colors_table.columns):
            Layout.set_column_width(column, self.COLORS_TABLE_WIDTHS[idx])

        # add the color description to the table in the corresponding cells
        self.add_color_description(colors_table, 0, 1, self.GREEN, 'No problem found')
        self.add_color_description(colors_table, 0, 5, self.ORANGE, 'Important problem found')
        self.add_color_description(colors_table, 2, 1, self.YELLOW, 'Marginal problem found')
        self.add_color_description(colors_table, 2, 5, self.RED, 'Critical problem found')

    def write_problem_description(self):
        """
        Write the description of the different problems.
        """

        for i in range(1, self.parameters['Number of problems'] + 1):
            self.report.add_paragraph(self.parameters['Problem {} description'.format(i)],
                                      'List Number 2')

    def write_chapter(self):
        """
        Write the whole chapter 'Effectiveness analysis', including the tables.
        """

        decision_table_index = self.tables.index(self.DECISION_TABLE)
        decision = DropDownLists.get_from_table(self.text_input_soup, decision_table_index)

        if decision[0] == 'Yes':
            effectiveness_analysis = ResultsChapter(self.report, self.text_input, self.text_input_soup, self.TITLE,
                                                    self.tables, self.picture_paths, self.parameters)

            self.report.add_paragraph(self.TITLE, self.TITLE_STYLE)
            self.make_result_table()
            self.make_colors_table()

            self.report.add_paragraph(self.DESCRIPTION_TITLE, self.DESCRIPTION_STYLE)
            self.write_problem_description()

            self.report.add_paragraph(self.DISCUSSION_TITLE, self.DISCUSSION_STYLE)
            effectiveness_analysis.write_chapter()
