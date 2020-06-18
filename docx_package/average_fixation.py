from docx.document import Document
from docx.table import Table
from bs4 import BeautifulSoup
from typing import List, Dict, Union
import seaborn as sns
import matplotlib.pyplot as plt
import pandas as pd
from docx.shared import Pt, Cm, RGBColor


from docx_package.results import ResultsChapter
from docx_package.picture import Picture
from docx_package import text_reading
from txt_package import plot, eye_tracking


class AverageFixation:
    """
    Class that represents the 'Average fixation' chapter and the visualization of its results.
    """

    # name of table as it appears in the tables list
    PLOT_TYPE_TABLE = 'Average fixation plot type table'

    # information about the headings of this chapter
    TITLE = 'Average fixation'
    TITLE_STYLE = 'Heading 2'
    DISCUSSION_TITLE = 'Discussion'
    DISCUSSION_STYLE = 'Heading 3'

    # path to plot image files
    PARTICIPANT_FIGURE_PATH = 'Outputs/Average_fixation_participant{}.png'
    BAR_PLOT_FIGURE_PATH = 'Outputs/Average_fixation_bar_plot.png'
    BOX_PLOT_FIGURE_PATH = 'Outputs/Average_fixation_box_plot.png'

    # caption of the pie plot figure
    BAR_PLOT_CAPTION = 'Bar plot showing the mean of the fixation time and the 95% confidence interval.'
    BOX_PLOT_CAPTION = 'Box plot showing the mean, the 25% and 75% quartiles and the distribution of the fixation time.'

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 text_input_soup: BeautifulSoup,
                 list_of_tables: List[str],
                 parameters_dictionary: Dict[str, Union[str, int]],
                 list_of_dataframes: List[pd.DataFrame]
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            list_of_tables: List of all table names.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value)
            list_of_dataframes: List of data frames containing the cGOM data of each participant
        """

        self.report = report_document
        self.text_input = text_input_document
        self.text_input_soup = text_input_soup
        self.tables = list_of_tables
        self.parameters = parameters_dictionary
        self.cGOM_dataframes = list_of_dataframes

    @ property
    def plot_type(self) -> str:
        """
        Returns:
            Dropdown list value of the parameter table corresponding to the plot type,
            i.e. 'Bar plot' or 'Box plot'.
        """

        plot_type_list = text_reading.get_dropdown_list_of_table(self.text_input_soup,
                                                                 self.tables.index(self.PLOT_TYPE_TABLE)
                                                                 )
        return plot_type_list[0]

    def make_plots(self):
        """
        Create plots to visualize the fixation durations.

        One box plot for each participant is created.
        One bar plot showing a confidence interval of 95% and one box plot with the data of
        all participants are created.
        """

        # main fixation times data frame
        average_fixation_df = pd.DataFrame()

        # create a data frame with the fixation times for each participant, create a box plot with it,
        # and append it to the main data frame
        for idx, dataframe in enumerate(self.cGOM_dataframes):
            aois = eye_tracking.areas_of_interest(dataframe)
            participant_fixations = eye_tracking.fixations(aois, dataframe)
            # TODO: choose which plot to make and complete docstring according to the choice
            plot.make_boxplot(data_frame=participant_fixations,
                              figure_save_path=self.PARTICIPANT_FIGURE_PATH.format(idx+1),
                              title='Average fixation: participant {}'.format(idx+1),
                              ylabel='Fixation time [s]')
            '''
            plot.make_barplot(data_frame=participant_fixations,
                              figure_save_path=self.PARTICIPANT_FIGURE_PATH.format(idx+1),
                              title='Average fixation: participant {}'.format(idx+1),
                              ylabel='Fixation time [s]')
            '''
            average_fixation_df = average_fixation_df.append(participant_fixations, ignore_index=True)

        # create a bar plot and a box plot with the fixations of all participants
        plot.make_boxplot(data_frame=average_fixation_df,
                          figure_save_path=self.BOX_PLOT_FIGURE_PATH,
                          title='Average fixation',
                          ylabel='Fixation time [s]')
        plot.make_barplot(data_frame=average_fixation_df,
                          figure_save_path=self.BAR_PLOT_FIGURE_PATH,
                          title='Average fixation',
                          ylabel='Fixation time [s]')

    def write_chapter(self):
        """
        Write the whole chapter 'Time on tasks', including the chosen plot.
        """

        self.make_plots()

        time_on_tasks = ResultsChapter(self.report, self.text_input, self.text_input_soup, self.TITLE,
                                       self.tables, self.parameters)

        self.report.add_paragraph(self.TITLE, self.TITLE_STYLE)

        # add bar plot or box plot depending on the choice of plot type
        if self.plot_type == 'Bar plot':
            Picture.add_picture_and_caption(self.report,
                                            [self.BAR_PLOT_FIGURE_PATH],
                                            self.BAR_PLOT_FIGURE_PATH,
                                            self.BAR_PLOT_CAPTION,
                                            width=Cm(12)
                                            )
        if self.plot_type == 'Box plot':
            Picture.add_picture_and_caption(self.report,
                                            [self.BOX_PLOT_FIGURE_PATH],
                                            self.BOX_PLOT_FIGURE_PATH,
                                            self.BOX_PLOT_CAPTION,
                                            width=Cm(12)
                                            )

        self.report.add_paragraph(self.DISCUSSION_TITLE, self.DISCUSSION_STYLE)
        time_on_tasks.write_chapter()
