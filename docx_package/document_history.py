from docx.document import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT


from docx_package.layout import Layout


class DocumentHistory:
    """
    Class that represents the chapter 'Document history'.
    """

    # information about the headings of this chapter
    TITLE = 'Document history'
    TITLE_STYLE = 'Heading 1'

    # entries of the first row
    FIRST_ROW = ['Version', 'Author', 'Description of changes']

    # hexadecimal of color for cell shading
    LIGHT_GREY_10 = 'D0CECE'

    # list of the table columns width
    TABLE_WIDTHS = [3, 5.75, 7.15]

    def __init__(self, report_document: Document):
        """
        Args:
            report_document: .docx file where the report is written.
        """

        self.report = report_document

    def add_table(self):
        """
        Add a table for the document history.
        """

        # create table and define its style
        history_table = self.report.add_table(rows=4, cols=3)
        history_table.style = 'Table Grid'
        history_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        history_table.autofit = True

        # fill the first row, color it in light_grey_10 and set the font to bold
        for idx, cell in enumerate(history_table.rows[0].cells):
            cell.text = self.FIRST_ROW[idx]
            Layout.set_cell_shading(cell, self.LIGHT_GREY_10)
            cell.paragraphs[0].runs[0].font.bold = True

        # set the vertical alignment of all cells
        for row in history_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # set the width of all columns
        for idx, column in enumerate(history_table.columns):
            Layout.set_column_width(column, self.TABLE_WIDTHS[idx])

    @ classmethod
    def write(cls, report_document: Document):
        """
        Write the chapter 'Document history' with its table.

        Args:
            report_document: .docx file where the report is written.
        """

        doc_history = cls(report_document)

        # add a heading to the chapter
        report_document.add_paragraph(doc_history.TITLE, doc_history.TITLE_STYLE)

        doc_history.add_table()
