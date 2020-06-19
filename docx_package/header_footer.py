from docx.section import Section
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from datetime import date
from typing import Dict, Union

from docx_package import layout


class Header:
    """
    Class that represents and creates the header of the report.
    """

    # header parameter keys
    FIRM_KEY = 'Firm name'
    HEADER_TITLE_KEY = 'Header title'
    VERSION_KEY = 'Version / ID'

    def __init__(self,
                 section: Section,
                 parameters_dictionary: Dict[str, Union[str, int]]
                 ):
        """
        Args:
            section: Report section to which we want to add the header.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value).
        """

        self.section = section
        self.parameters = parameters_dictionary

    @ staticmethod
    def add_tab_stops(paragraph: Paragraph):
        """
        Add three tab stops to a paragraph, i.e. left, center and right.

        Args:
            paragraph: Paragraph where the tab stops will be added.
        """

        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(0), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.CENTER, WD_TAB_LEADER.SPACES)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

    def write(self):
        """
        Write a header in the section of the report.

        The first line of the header contains the firm name on the left, the title in the middle
        and the version number on the right.
        The second line of the header contains the date on the right.
        """

        header = self.section.header

        # header of both sections of the report are not linked
        header.is_linked_to_previous = False

        # create two lines of header with tab stops
        first_line = header.paragraphs[0]
        second_line = header.add_paragraph()
        first_line.paragraph_format.tab_stops.clear_all()
        self.add_tab_stops(first_line)
        self.add_tab_stops(second_line)

        # create the entries of the header
        firm = layout.capitalize_first_letter(self.parameters[self.FIRM_KEY])
        title = layout.capitalize_first_letter(self.parameters[self.HEADER_TITLE_KEY])
        version = layout.capitalize_first_letter(self.parameters[self.VERSION_KEY])
        today_date = date.today()
        date_string = today_date.strftime('%d.%m.%Y')

        # write the entries in the header
        first_line.text = '{} \t {} \t {}'.format(firm, title, version)
        second_line.text = ' \t \t {}'.format(date_string)


class Footer:
    def __init__(self,
                 section: Section
                 ):
        """
        Args:
            section: Report section to which we want to add the header.
        """

        self.section = section

    @ staticmethod
    def add_tab_stops(paragraph: Paragraph):
        """
        Add three tab stops to a paragraph, i.e. left, center and right.

        Args:
            paragraph: Paragraph where the tab stops will be added.
        """

        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(0), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.CENTER, WD_TAB_LEADER.SPACES)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

    @ staticmethod
    def add_page_number(run: Run):
        """
        Add the page number to a paragraph.

        Args:
            run: Run in which the page number will be added.
        """

        # add XML elements and set their attributes so that the page number correspond to a real page number
        r_element = run._r

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r_element.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        r_element.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r_element.append(fldChar2)

    def write(self):
        """
        Write a footer in the section of the report.

        The second line of the footer contains the page number on the right.
        """

        footer = self.section.footer

        # footer of both sections of the report are not linked
        footer.is_linked_to_previous = False

        # create two lines of footer with tab stops
        first_line = footer.paragraphs[0]
        second_line = footer.add_paragraph()
        first_line.paragraph_format.tab_stops.clear_all()
        self.add_tab_stops(first_line)
        self.add_tab_stops(second_line)

        # add the page number on the right-hand side of the footer
        second_line.text = ' \t \t'
        self.add_page_number(second_line.add_run())
