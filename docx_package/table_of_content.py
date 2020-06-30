from docx.document import Document
from docx.shared import Cm
from typing import List, Dict, Union
from bs4 import BeautifulSoup

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import win32com.client
import inspect, os


class TableOfContent:
    def __init__(self, report_document: Document):
        """
        Args:
            report_document: .docx file where the report is written.
        """

        self.report = report_document

    def write(self):
        """
        Add a table of content.

        The table of content will not appear at the first time.
        It has to be updated by pressing Ctrl + A, and then F9.

        Args:
            report_document: .docx file where the report is written.
        """

        # add the heading of the table of content
        self.report.add_paragraph('Table of content', 'Table of content')

        # access to XML run element <w:r>
        paragraph = self.report.add_paragraph()
        run = paragraph.add_run()
        r = run._r

        # create new XML elements, set their attributes and add them to the run element
        # so that the table of content is considered as such and can be updated
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-2" \\h \\z \\u'  # change 1-3 depending on heading levels you need
        r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = 'Press "Ctrl + A" to select everything and then "F9" to update fields.'
        fldChar2.append(fldChar3)
        r.append(fldChar2)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        r.append(fldChar4)

    def update(self, report_file_path):
        word = win32com.client.DispatchEx("Word.Application")
        doc = word.Documents.Open(report_file_path)

        word.ActiveDocument.Fields.Update()

        # doc.TablesOfContents(1).Update()
        doc.Close(SaveChanges=True)
        word.Quit()
