from docx.document import Document
from bs4 import BeautifulSoup
from typing import List

from docx_package.dropdown_lists import DropDownLists


class Parameters:
    """
    Class that represents and defines the parameter used to write the chapters of the report.
    """

    # tables with two columns having parameters in each row
    # all those tables are handled the same way
    STANDARD_PARAMETERS_TABLES = [
        'Study table',
        'Title table',
        'Header table',
        'Approval table',
    ]

    CHARACTERISTICS_TABLE = 'Participants characteristics table'

    # tables that have special features (e.g. more than two columns, dropdown lists, ...)
    # those tables are handled separately
    TASKS_TABLE = 'Critical tasks description table'
    PROBLEMS_TABLE = 'Effectiveness analysis problem type table'

    def __init__(self,
                 text_input_document: Document,
                 text_input_soup: BeautifulSoup,
                 list_of_tables: List[str]
                 ):
        """
        Args:
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            list_of_tables: List of all table names.
        """

        self.text_input = text_input_document
        self.text_input_soup = text_input_soup
        self.tables = list_of_tables

        # dictionary where keys and values of all parameters will be stored
        self.dictionary = {}

    def get_from_standard_tables(self):
        """
        Read the parameters from the standard tables and stored them in the dictionary.

        Standard tables have two columns, the first one contains the keys and the second one the values of parameters.
        """

        for table_name in self.STANDARD_PARAMETERS_TABLES:
            table_index = self.tables.index(table_name)
            table = self.text_input.tables[table_index]

            for row in table.rows:
                key = row.cells[0].text
                value_text = row.cells[1].text

                # TODO: add error message if given number does not correspond to number of element

                # case where the value is an integer
                if key.startswith('Number of'):

                    if value_text and value_text.isdigit() and int(value_text) <= 15:
                        self.dictionary[key] = int(value_text)

                    # if no number where provided,
                    # read it from the number of described elements in the corresponding table
                    else:
                        if 'participants' in key:
                            self.dictionary[key] = self.get_number(self.tables.index(self.CHARACTERISTICS_TABLE))
                        if 'tasks' in key:
                            self.dictionary[key] = self.get_number(self.tables.index(self.TASKS_TABLE))

                # case where the value is a string
                else:
                    self.dictionary[key] = value_text

    def get_number(self, table_index: int) -> int:
        """
        Get the number of described elements in a table.

        This function is called to determine the number of participants or the number of critical tasks
        if they were not provided in the text input form.

        Args:
            table_index: Index of the table where the elements, i.e. participants or critical tasks, are described.

        Returns:
            The number of described elements, i.e. number of participants or number of critical tasks.
        """

        table = self.text_input.tables[table_index]

        # return the index of a row when nothing was written in it
        for idx, row in enumerate(table.rows[1:]):
            row_described = False
            for cell in row.cells[1:]:
                if cell.text:
                    row_described = True
            if not row_described:
                return idx

    def get_from_tasks_table(self):
        """
        Read the parameters from the critical tasks table and stored them in the dictionary.

        The critical tasks table differs from the standard table because it has 3 columns.
        """

        tasks_table_index = self.tables.index(self.TASKS_TABLE)
        tasks_table = self.text_input.tables[tasks_table_index]

        for row in tasks_table.rows[1:]:
            type_key = row.cells[0].text + ' name'
            description_key = row.cells[0].text + ' description'
            type_value = row.cells[1].text
            description_value = row.cells[2].text

            # add keys and values to dictionary if a critical is defined either with its type or its definition
            if type_value.replace(' ', '') or description_value.replace(' ', ''):
                self.dictionary[type_key] = type_value
                self.dictionary[description_key] = description_value

    def get_from_problems_table(self):
        """
        Read the parameters from the problems table and stored them in the dictionary.

        The problems table differs from the standard table because it has 3 columns and contains dropdown lists.
        """

        problems_table_index = self.tables.index(self.PROBLEMS_TABLE)
        problems_table = self.text_input.tables[problems_table_index]
        problem_types = DropDownLists.get_from_table(self.text_input_soup, problems_table_index)

        # return the index of a row when nothing was written in it to get the number of problems
        problems_number = 0
        for idx, problem_type in enumerate(problem_types):
            if '-' in problem_type:
                problems_number = idx
                break

        self.dictionary['Number of problems'] = problems_number

        # cells are not recognized as cells by word if they contain a dropdown list,
        # so here is a work around to get the values

        # list of the text of all cells, except those ones containing dropdown list
        list_of_text = []

        # stores the text of every relevant cell in the list
        if problems_number != 0:
            stop = False
            while not stop:
                for i in range(1, problems_number + 1):
                    for cell in problems_table.rows[i].cells:
                        text = cell.text
                        if text or len(list_of_text) == (problems_number * 2) - 1:
                            list_of_text.append(text)
                        else:
                            list_of_text.pop()
                            stop = True

        # delete the last item to ensure that there is no key without a corresponding value
        if len(list_of_text) % 2 != 0:
            list_of_text.pop()

        value = 0

        # stores the keys and corresponding values in the dictionary
        for i in range(len(list_of_text)):

            # the problem number appears every two items
            if i % 2 == 0:

                # problem types
                problem_number = list_of_text[i]
                type_key = problem_number + ' type'
                self.dictionary[type_key] = problem_types[value]

                # problem descriptions
                description = list_of_text[i + 1]
                description_key = problem_number + ' description'
                self.dictionary[description_key] = description

                value += 1

    @ classmethod
    def get_all(cls, text_input_document, text_input_soup, list_of_tables):
        """
        Args:
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            list_of_tables: List of all table names.

        Returns:
            Dictionary containing values and keys of all parameters.
        """

        parameters = cls(text_input_document, text_input_soup, list_of_tables)

        parameters.get_from_standard_tables()
        parameters.get_from_tasks_table()
        parameters.get_from_problems_table()

        return parameters.dictionary
