from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Cm
from bs4 import BeautifulSoup
from typing import List, Dict, Union
import numpy as np
import pandas as pd

from docx_package.layout import Layout
from docx_package.results import ResultsChapter
from docx_package.picture import Picture
from eye_tracking_package import plot, eye_tracking
from docx_package import text_reading


class DwellTimesAndRevisits:
    """
    Class that represents the 'Dwells times and revisits' chapter and the visualization of its results.
    """

    # name of tables as they appear in the tables list
    DECISION_TABLE = 'Dwell times and revisits decision table'

    # information about the headings of this chapter
    TITLE = 'Dwell times and revisits'
    TITLE_STYLE = 'Heading 2'
    DISCUSSION_TITLE = 'Discussion'
    DISCUSSION_STYLE = 'Heading 3'

    # index name of the statistics
    SUM_INDEX = 'Sum'
    MEAN_INDEX = 'Mean'
    MAX_INDEX = 'Max'
    MIN_INDEX = 'Min'

    # text of the cells of the first row
    TABLE_FIRST_ROW = ['Areas of interest', 'Dwell times [s]', 'Average [s]', 'Max [s]', 'Min [s]', 'Revisits']

    # color for cell shading
    LIGHT_GREY_10 = 'D0CECE'

    # list of the table column width
    WIDTHS = [4, 2.8, 2.29, 2.29, 2.29, 2.29]

    # path to plot image files
    PARTICIPANT_FIGURE_PATH = 'Outputs/Dwell_times_participant{}.png'
    PIE_PLOT_FIGURE_PATH = 'Outputs/Dwell_times_pie_plot.png'

    # caption of the pie plot figure
    CAPTION = 'Average total dwell times amount for each area of interest over all participants.'

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 text_input_soup: BeautifulSoup,
                 list_of_tables: List[str],
                 picture_paths_list: List[str],
                 parameters_dictionary: Dict[str, Union[str, int]],
                 list_of_dataframes: List[pd.DataFrame]
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            list_of_tables: List of all table names.
            picture_paths_list: List of the path of all input pictures.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value)
            list_of_dataframes: List of data frames containing the cGOM data of each participant
        """

        self.report = report_document
        self.text_input = text_input_document
        self.text_input_soup = text_input_soup
        self.tables = list_of_tables
        self.picture_paths = picture_paths_list
        self.parameters = parameters_dictionary
        self.cGOM_dataframes = list_of_dataframes

    def make_dwell_times_plot_and_dataframe(self) -> pd.DataFrame:
        """
        Create pie plots of the total sums of dwell times.

        One pie plot for each participant is created.
        One pie plot with the data of all participants is created.

        Returns:
            Data frame with the average of the statistics for each participants and each AOI.

            The indexes are the AOIs and the columns are the statistics,
            i.e. 'Sum', 'Mean', 'Max', 'Min'.
        """

        # main dwell times data frame that contains the dwell times of all participants
        all_dwell_times_df = pd.DataFrame()

        # create a data frame with the dwell times and statistics of each participants
        # and append it to the main data frame
        for idx, dataframe in enumerate(self.cGOM_dataframes):
            aois = eye_tracking.areas_of_interest(dataframe)
            participants_df = eye_tracking.dwell_times(aois, dataframe)
            all_dwell_times_df = all_dwell_times_df.append(participants_df)

            # plot the total sum of the dwell times for each participants
            participant_sum = participants_df[self.SUM_INDEX].to_numpy()
            participant_sum = participant_sum[~np.isnan(participant_sum)]
            plot.make_pieplot(data_vector=participant_sum,
                              labels_list=aois,
                              figure_save_path=self.PARTICIPANT_FIGURE_PATH.format(idx + 1),
                              title='Dwell times: participant {}'.format(idx + 1)
                              )


        # create a data frame with the mean of the statistics for all participants for each AOI
        all_aois = eye_tracking.areas_of_interest(all_dwell_times_df)
        dwell_times_table = pd.DataFrame(index=all_aois,
                                         columns=[self.SUM_INDEX, self.MEAN_INDEX, self.MAX_INDEX, self.MIN_INDEX],
                                         data=np.zeros((len(all_aois), 4))
                                         )

        # calculate the mean of all statistics and store it in the data frame
        for aoi in all_aois:
            data_of_aoi = all_dwell_times_df[all_dwell_times_df.index == aoi]
            dwell_times_table.loc[aoi].at[self.SUM_INDEX] = data_of_aoi[self.SUM_INDEX].mean()
            dwell_times_table.loc[aoi].at[self.MEAN_INDEX] = data_of_aoi[self.MEAN_INDEX].mean()
            dwell_times_table.loc[aoi].at[self.MAX_INDEX] = data_of_aoi[self.MAX_INDEX].mean()
            dwell_times_table.loc[aoi].at[self.MIN_INDEX] = data_of_aoi[self.MIN_INDEX].mean()

        # create a pie plot with the average dwell times sum of all participants or
        # do nothing if no cGOM data is provided
        all_sums = dwell_times_table[self.SUM_INDEX].to_numpy()
        if not dwell_times_table.empty:
            plot.make_pieplot(data_vector=all_sums,
                              labels_list=all_aois,
                              figure_save_path=self.PIE_PLOT_FIGURE_PATH,
                              title='Dwell times'
                              )

        return dwell_times_table

    def revisits_stat(self) -> pd.DataFrame:
        """
        Returns:
            Data frame containing the revisits of all participants (index) and all tasks (columns).
            The last row contains the mean revisits for each task.
        """

        # main revisits data frame
        revisits_df = pd.DataFrame()

        # create a data frame with the revisits for each participant and append it to the main data frame
        for idx, dataframe in enumerate(self.cGOM_dataframes):
            aois = eye_tracking.areas_of_interest(dataframe)
            revisits = eye_tracking.revisits(aois, dataframe)
            participant_revisits = pd.DataFrame(index=['Participant {}'.format(idx + 1)],
                                                columns=aois,
                                                data=[revisits]
                                                )
            revisits_df = revisits_df.append(participant_revisits)

        # calculate the mean of revisits for each AOI and append it to the main data frame
        revisits_mean = revisits_df.mean().to_numpy()
        revisits_mean_df = pd.DataFrame(index=[self.MEAN_INDEX],
                                        columns=revisits_df.columns,
                                        data=[revisits_mean]
                                        )
        revisits_df = revisits_df.append(revisits_mean_df)

        return revisits_df

    def add_table(self):
        """
        Create the table with the statistics of the dwell times and revisits,
        i.e. the total dwell times amount, the average dwell time,
        the maximum and minimum dwell time and the revisits.
        """

        # data frames containing the data
        dwell_times_df = self.make_dwell_times_plot_and_dataframe()
        revisits_df = self.revisits_stat()

        # do nothing if no cGOM data is provided
        if not dwell_times_df.empty:
            # areas of interests
            aois = dwell_times_df.index

            # create table
            table = self.report.add_table(len(aois) + 1, len(self.TABLE_FIRST_ROW))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            # write the first row
            for index, label in enumerate(self.TABLE_FIRST_ROW):
                cell = table.rows[0].cells[index]
                cell.text = label
                Layout.set_cell_shading(cell, self.LIGHT_GREY_10)  # color the cell in light_grey_10
                cell.paragraphs[0].runs[0].font.bold = True

            # write the first column with the name of the areas of interest
            for idx, aoi in enumerate(aois):
                cell = table.columns[0].cells[idx + 1]
                cell.text = aoi

            # write all the entries of the tables except the ones about revisits
            matrix = dwell_times_df.to_numpy()
            for i in range(len(aois)):
                for j in range(4):
                    table.cell(i + 1, j + 1).text = str(round(matrix[i, j], 4))

            # write the third column with the mean of revisits of each AOI
            for idx, revisits in enumerate(revisits_df.loc[self.MEAN_INDEX].to_numpy()):
                cell = table.columns[5].cells[idx + 1]
                cell.text = str(round(revisits, 4))

            # set the vertical and horizontal alignment of all cells
            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for i in range(len(table.rows)):
                for j in range(1, len(table.columns)):
                    table.cell(i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # set the widths of all columns
            for idx, column in enumerate(table.columns):
                Layout.set_column_width(column, self.WIDTHS[idx])

    def write_chapter(self):
        """
        Write the whole chapter 'Dwell times and revisits', including the table and the plot.
        """

        decision_table_index = self.tables.index(self.DECISION_TABLE)
        decision = text_reading.get_dropdown_list_of_table(self.text_input_soup, decision_table_index)

        if decision[0] == 'Yes':
            time_on_tasks = ResultsChapter(self.report, self.text_input, self.text_input_soup, self.TITLE,
                                           self.tables, self.picture_paths, self.parameters)

            self.report.add_paragraph(self.TITLE, self.TITLE_STYLE)

            self.add_table()

            try:
                Picture.add_picture_and_caption(self.report,
                                                [self.PIE_PLOT_FIGURE_PATH],
                                                self.PIE_PLOT_FIGURE_PATH,
                                                self.CAPTION,
                                                width=Cm(12)
                                                )
            except FileNotFoundError:     # do nothing if no cGOM data is provided
                pass

            self.report.add_paragraph(self.DISCUSSION_TITLE, self.DISCUSSION_STYLE)
            time_on_tasks.write_chapter()
