from typing import List, Union
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from os import listdir
from PIL import Image, UnidentifiedImageError


class Picture:
    """
    Class that represents everything that have something to do with the pictures in the report,
    i.e. creating list of picture paths, adding pictures, adding captions,
    adding list of figures or printing error message regarding pictures.

    Pictures must be saved in the 'Inputs/Pictures' directory in image format (e.g. .jpg, .jpeg, .png, .gif, ...).
    """

    # information for the caption
    CAPTION_LABEL = 'Figure '
    CAPTION_STYLE = 'Caption'

    def __init__(self,
                 report_document: Document,
                 picture_paths: List[str],
                 picture_name: str,
                 caption_text: str,
                 picture_width: Union[Cm, None],
                 picture_height: Union[Cm, None],
                 space_before: Union[Cm, None],
                 space_after: Union[Cm, None]
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            picture_paths: List of paths of all input pictures.
            picture_name: Name of the picture file without the extension.
            caption_text: Text of the picture caption.
            picture_width: Width of the picture as it appears in the report.
            picture_height: Height of the picture as it appears in the report.
            space_before: Space before the paragraph of the picture.
            space_after: Space after the paragraph of the caption.
        """

        self.report = report_document
        self.picture_paths = picture_paths
        self.picture_name = picture_name
        self.caption = caption_text
        self.width = picture_width
        self.height = picture_height
        self.space_before = space_before
        self.space_after = space_after

    @ staticmethod
    def get_picture_paths() -> List[str]:
        """
        Returns:
            List of paths of all pictures given as input.
        """

        # list of names of all files stored in the directory 'Inputs/Pictures'
        pictures = listdir('Inputs/Pictures')

        # create a list of paths to all files
        picture_paths = []
        for picture in pictures:
            path = 'Inputs/Pictures/{}'.format(picture)
            picture_paths.append(path)

        return picture_paths

    def add_picture(self) -> bool:
        """
        Add a picture to the report, that is centered w.r.t. the margin.

        Returns:
            True if a picture was added, and False if not.
        """

        # find the files that correspond to the picture file name
        for index, picture_path in enumerate(self.picture_paths):
            if self.picture_name in picture_path:

                try:
                    # try to open the file to control that it is an image
                    Image.open(picture_path)

                    # add a picture with the given size in the center of the side margin
                    picture_paragraph = self.report.add_paragraph(style='Picture')
                    picture_paragraph.add_run().add_picture(picture_path, width=self.width, height=self.height)

                    # set space before the paragraph of the picture
                    picture_paragraph.paragraph_format.space_before = self.space_before

                    # delete the path from the list because this picture will not be added again
                    self.picture_paths.pop(index)

                    # terminate because a picture was added and return True
                    return True

                # do nothing if the file is not an image
                except UnidentifiedImageError:
                    pass

        # return False because no picture was added
        return False

    def add_caption(self):
        """
        Add a caption of the form: 'Figure <figure number>: <caption text>, e.g. 'Figure 3: A medical device.'

        The caption will not appear in this form at the first time.
        It has to be updated by pressing Ctrl + A, and then F9.
        """

        # add the label of the caption
        caption_paragraph = self.report.add_paragraph(self.CAPTION_LABEL, style=self.CAPTION_STYLE)

        # add XML elements and set their attributes so that the caption is considered as such and can be updated
        run = caption_paragraph.add_run()
        r_element = run._r

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r_element.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.text = 'SEQ Figure \\* ARABIC'
        r_element.append(instrText)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        r_element.append(fldChar)

        # add the text of the caption
        caption_paragraph.add_run(': {}'.format(self.caption))

        # set space after the paragraph of the caption
        caption_paragraph.paragraph_format.space_after = self.space_after

    @ classmethod
    def add_picture_and_caption(cls,
                                report_document: Document,
                                picture_paths: List[str],
                                picture_name: str,
                                caption: str,
                                width=None,
                                height=None,
                                space_before=None,
                                space_after=None
                                ):
        """
        Add a picture to the report if there is one that corresponds to the picture file name
        and a caption after the picture if one was added.

        Args:
            report_document: .docx file where the report is written.
            picture_paths: List of paths of all input pictures.
            picture_name: Name of the picture file without the extension.
            caption: Text of the picture caption.
            width (optional): Width of the picture as it appears in the report.
            height (optional): Height of the picture as it appears in the report.
            space_before (optional): Space before the paragraph of the picture.
                                     None if inherited from the style hierarchy.
            space_after (optional): Space after the paragraph of the caption.
                                    None if inherited from the style hierarchy.
        """

        picture = cls(report_document, picture_paths, picture_name, caption, width, height, space_before, space_after)
        picture_added = picture.add_picture()
        if picture_added:
            picture.add_caption()

    @ staticmethod
    def add_figures_list(report_document):
        """
        Add a list of figures.

        The list of figures will not appear at the first time.
        It has to be updated by pressing Ctrl + A, and then F9.

        Args:
            report_document: .docx file where the report is written.
        """

        # add the heading of the list of figures
        report_document.add_paragraph('List of figures', 'Heading 2')

        # access to XML run element <w:r>
        paragraph = report_document.add_paragraph()
        run = paragraph.add_run()
        r = run._r

        # create new XML elements, set their attributes and add them to the run element
        # so that the list of figure is considered as such and can be updated
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\h \\z \\c \"Figure\"'
        r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = 'Press "Ctrl + A" to select everything and then "F9" to update fields.'
        fldChar2.append(fldChar3)
        r.append(fldChar2)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        r.append(fldChar4)

    @ staticmethod
    def error_message(picture_paths):
        """
        Print an error message that show the path of the pictures that were not added to the report
        and give some possible problems that might have occurred.

        Args:
            picture_paths: List of paths of all remaining input pictures.
        """

        print('These pictures were not added to the report:')

        # print path of all pictures that were added
        for picture_path in picture_paths:
            print('   ', picture_path)

        print('\nPossible problems are: \n',
              '   1. It is not an image file. \n',
              '   2. More than 3 pictures were given for a chapter. \n',
              '   3. The name of the file is unexpected.'
              )
        print('\n----------\n')
