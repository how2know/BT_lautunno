from docx.document import Document
from docx.table import Table, _Row
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from typing import List, Dict, Union, Tuple

from docx_package.layout import Layout


class ParticipantsCharacteristics:
    """
    Class that represents the chapter 'Participants’ characteristics' in the appendix of the report.
    """

    # information about the headings of this chapter
    TITLE = 'Participants’ characteristics'
    TITLE_STYLE = 'Heading 2'

    # participants' characteristics table as it appears in the table list
    CHARACTERISTICS_TABLE = 'Participants characteristics table'

    # participants number parameter key
    PARTICIPANTS_NUMBER_KEY = 'Number of participants'

    # hexadecimal of color for cell shading
    LIGHT_GREY_10 = 'D0CECE'

    # list of the table columns width
    TABLE_WIDTHS = [2.3, 1.6, 1, 2.9, 2.2, 3, 2.9]

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 list_of_tables: List[str],
                 parameters_dictionary: Dict[str, Union[str, int]]
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            list_of_tables: List of all table names.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value)
        """

        self.report = report_document
        self.text_input = text_input_document
        self.tables = list_of_tables
        self.parameters = parameters_dictionary

    @ property
    def input_table(self) -> Table:
        """
        Returns:
            Table of the input .docx file where the participants' characteristics are written.
        """

        participants_table_index = self.tables.index(self.CHARACTERISTICS_TABLE)
        return self.text_input.tables[participants_table_index]

    @ property
    def described_rows(self) -> Tuple[int, List[_Row]]:
        """
        Get the number of described elements in a table.

        Returns:
            Tuple that contains the number of participants described in the participant characteristics table
             and a list of rows in which they are described.
        """

        participants_number = 0
        described_rows = []

        # store the row in the list if it is described and increase the number of participants
        for row in self.input_table.rows[1:]:
            row_described = False
            for cell in row.cells[1:]:
                if cell.text:
                    row_described = True
            if row_described:
                participants_number += 1
                described_rows.append(row)

        return participants_number, described_rows

    def add_table(self):
        """
        Add a table for the document history.
        """

        participants_number = self.described_rows[0]
        described_rows = self.described_rows[1]

        # create table and define its style
        rows_number = participants_number + 1
        cols_number = len(self.input_table.columns)
        appendix_table = self.report.add_table(rows_number, cols_number)
        appendix_table.style = 'Table Grid'
        appendix_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        appendix_table.autofit = True

        for i in range(rows_number):
            for j in range(cols_number):

                # fill the first row
                if i == 0:
                    appendix_table.cell(i, j).text = self.input_table.cell(i, j).text

                # fill the first columns with 'P1', 'P2', 'P3', etc... with the number corresponding to the participant
                elif i != 0 and j == 0:
                    appendix_table.cell(i, j).text = 'P{}'.format(i)

                # fill all other cells with the entries given in the described rows
                else:
                    appendix_table.cell(i, j).text = described_rows[i-1].cells[j].text

        # color the first row in light_grey_10 and set the font to bold
        for cell in appendix_table.rows[0].cells:
            Layout.set_cell_shading(cell, self.LIGHT_GREY_10)
            cell.paragraphs[0].runs[0].font.bold = True

        # set the vertical and horizontal alignment of all cells
        for row in appendix_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for i in range(rows_number):
            for j in range(cols_number):
                if i == 0 or j == 0:
                    appendix_table.cell(i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # set the width of all columns
        for idx, column in enumerate(appendix_table.columns):
            Layout.set_column_width(column, self.TABLE_WIDTHS[idx])

    @ classmethod
    def write(cls,
              report_document: Document,
              text_input_document: Document,
              list_of_tables: List[str],
              parameters_dictionary: Dict[str, Union[str, int]]
              ):
        """
        Write the chapter 'Participants’ characteristics' with its table if some participants were described.

        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            list_of_tables: List of all table names.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value)
        """

        participant_appendix = cls(report_document, text_input_document, list_of_tables, parameters_dictionary)

        if participant_appendix.described_rows[0] != 0:

            # add a heading to the chapter
            report_document.add_paragraph(participant_appendix.TITLE, participant_appendix.TITLE_STYLE)

            participant_appendix.add_table()
