from docx import Document
from typing import List
from bs4 import BeautifulSoup

from docx_package import text_reading


class Definitions:
    """
    Class that represents the 'Definitions' chapter.
    """

    # information about the title of this chapter
    TITLE = 'Terms definitions'
    TITLE_STYLE = 'Heading 1'

    # information about the title of this chapter
    REFERENCES_TITLE = 'References'
    REFERENCES_TITLE_STYLE = 'Heading 2'

    # name of the standards as they appear in the definitions document
    STANDARDS_NAMES = ['EU Regulation 2017/745', 'IEC 62366-1', 'FDA Guidance']

    # dictionary where the keys are the defined terms and the values are a tuple of
    # a list of paragraphs corresponding to the definitions and
    # a list of the styles in which these paragraphs will be written
    DEFINITIONS_DICTIONARY = {}

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 text_input_soup: BeautifulSoup,
                 definitions_document: Document,
                 list_of_tables: List[str],
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            definitions_document: .docx file where all definitions are written.
            list_of_tables: List of all table names.
        """

        self.report = report_document
        self.text_input = text_input_document
        self.text_input_soup = text_input_soup
        self.definitions = definitions_document
        self.list_of_tables = list_of_tables

    def standard_heading_index(self, standard_name: str) -> int:
        """
        Args:
            standard_name: Name of the standard of which we want the index.

        Returns:
            Paragraph index of the chapter heading of the standard in the definitions document.
        """

        for paragraph_index, paragraph in enumerate(self.definitions.paragraphs):
            if paragraph.text == standard_name and 'Heading' in paragraph.style.name:
                return paragraph_index

    def next_standard_heading_index(self, previous_index) -> int:
        for paragraph_index, paragraph in enumerate(self.definitions.paragraphs[previous_index + 1:]):
            if paragraph.text in self.STANDARDS_NAMES and 'Heading' in paragraph.style.name:
                return paragraph_index + previous_index + 1

    def standard_wanted_terms(self, standard_name: str) -> List[str]:
        """
        Args:
            standard_name: Name of the standard of which we want the terms that have to be defined.

        Returns:
            List of all terms that have to be defined for this standard.
        """

        # select the table that have the information for this standard
        for table_index, table in enumerate(self.list_of_tables):
            if standard_name in table:

                # create a list of 'Yes' or 'No' stored in dropdown lists in the table of this standard
                list_of_yes_no = text_reading.get_dropdown_list_of_table(self.text_input_soup, table_index)

                # create a list of all possible terms that can be defined in this standard
                list_of_terms = []
                for row in self.text_input.tables[table_index].rows:
                    for cell in row.cells:
                        if cell.text:
                            list_of_terms.append(cell.text)

                # create a list of all terms that have to be defined for this standard
                list_of_defined_terms = []
                for index, term in enumerate(list_of_terms):
                    if list_of_yes_no[index] == 'Yes':
                        list_of_defined_terms.append(term)

                return list_of_defined_terms

    '''
    def write_terms_definitions(self, standard_name: str):
        """
        Write in the report the terms and their definition given in this standard.

        Args:
            standard_name: Name of the standard of which we want to write the definitions.
        """

        # create a list of paragraph indexes of all headings in the standard section,
        # i.e. the indexes of the terms
        terms_heading_indexes = []
        standard_heading_index = self.standard_heading_index(standard_name)
        next_index = self.next_standard_heading_index(standard_heading_index)
        for paragraph_index, paragraph in enumerate(self.definitions.paragraphs[standard_heading_index: next_index]):
            if 'Heading' in paragraph.style.name:
                terms_heading_indexes.append(paragraph_index + standard_heading_index)

        # create a list of all paragraphs that have to be written, i.e. the terms and their definitions,
        # and a list of the style in which the paragraphs have to be written
        list_of_paragraphs = []
        list_of_styles = []
        wanted_terms = self.standard_wanted_terms(standard_name)
        for index, terms_heading_index in enumerate(terms_heading_indexes):
            if self.definitions.paragraphs[terms_heading_index].text in wanted_terms:
                for i in range(terms_heading_index, terms_heading_indexes[index + 1]):
                    list_of_paragraphs.append(self.definitions.paragraphs[i].text)
                    list_of_styles.append(self.definitions.paragraphs[i].style.name)

        # write the paragraphs, i.e. the terms and their definitions, with the according style in the report
        for index, paragraph in enumerate(list_of_paragraphs):
            self.report.add_paragraph(paragraph, list_of_styles[index])
    '''

    def store_definitions(self, standard_name: str, reference_number: int):
        """
        Store the terms that have to defined in the dictionary.

        The keys corresponds to the name of the terms that have to be defined.
        The values are a tuple containing a list of the paragraphs that corresponds to the term and its definitions
        and a list of styles in which the paragraph must be written.

        Args:
            standard_name: Name of the standard of which we want to write the definitions.
            reference_number: Number that corresponds to the reference of the standard in the report.
        """

        # create a list of paragraph indexes of all headings in the standard section,
        # i.e. the indexes of the terms
        terms_heading_indexes = []
        standard_heading_index = self.standard_heading_index(standard_name)
        next_index = self.next_standard_heading_index(standard_heading_index)
        for idx, paragraph in enumerate(self.definitions.paragraphs[standard_heading_index: next_index]):
            if 'Heading' in paragraph.style.name:
                terms_heading_indexes.append(idx + standard_heading_index)

        # select the indexes of the terms that have to be defined
        wanted_terms = self.standard_wanted_terms(standard_name)
        for idx, terms_heading_index in enumerate(terms_heading_indexes):
            term = self.definitions.paragraphs[terms_heading_index].text
            if term in wanted_terms:

                # create a list of all paragraphs that have to be written, i.e. the terms and their definitions,
                # and a list of the style in which the paragraphs have to be written
                list_of_paragraphs = []
                list_of_styles = []
                for i in range(terms_heading_index, terms_heading_indexes[idx + 1]):
                    list_of_paragraphs.append(self.definitions.paragraphs[i].text)
                    list_of_styles.append(self.definitions.paragraphs[i].style.name)

                # add the references at the end of the paragraph
                list_of_paragraphs.append(' [{}]'.format(reference_number))

                # store the defined term as key in the dictionary,
                # and a tuple containing the list of paragraphs and the list of styles as value
                if term not in self.DEFINITIONS_DICTIONARY.keys():
                    self.DEFINITIONS_DICTIONARY[term] = (list_of_paragraphs, list_of_styles)

                # if the defined term is already stored as key from another standard,
                # add a space to its name to store it
                else:
                    self.DEFINITIONS_DICTIONARY[term + ' '] = (list_of_paragraphs, list_of_styles)

    def write_definitions(self):
        """
        Write the terms of all standards and their definition in the report.
        """
        ref_number = 1

        # store all terms that have to be defined in the dictionary
        for idx, standard_name in enumerate(self.STANDARDS_NAMES):
            self.store_definitions(standard_name, ref_number)

            # use the next reference number for the next standard
            if self.standard_wanted_terms(standard_name):
                ref_number += 1

        # sort all the terms alphabetically
        sorted_terms = sorted(self.DEFINITIONS_DICTIONARY.keys())

        # write the terms and their definitions in the report
        for term in sorted_terms:
            paragraphs = self.DEFINITIONS_DICTIONARY[term][0]
            styles = self.DEFINITIONS_DICTIONARY[term][1]
            for idx, paragraph in enumerate(paragraphs[:-1]):
                par = self.report.add_paragraph(paragraph, styles[idx])

            # add the reference at the end of the last paragraph
            par.add_run(paragraphs[-1])

    @ property
    def references(self):
        """
        Returns:
            List of the references to all standard.

        """

        references_heading_index = self.standard_heading_index('References')
        references = []

        for paragraph in self.definitions.paragraphs[references_heading_index+1:]:
            references.append(paragraph.text)

        return references

    @ classmethod
    def write_references(cls,
                         report_document: Document,
                         text_input_document: Document,
                         text_input_soup: BeautifulSoup,
                         definitions_document: Document,
                         list_of_tables: List[str]
                         ):
        """
        Write the 'References' chapter

        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            definitions_document: .docx file where all definitions are written.
            list_of_tables: List of all table names.
        """

        definitions = cls(report_document, text_input_document, text_input_soup, definitions_document, list_of_tables)

        report_document.add_paragraph(definitions.REFERENCES_TITLE, definitions.REFERENCES_TITLE_STYLE)

        # write the references to the standard that were used to defined the terms
        for idx, standard_name in enumerate(definitions.STANDARDS_NAMES):
            if definitions.standard_wanted_terms(standard_name):
                definitions.report.add_paragraph(definitions.references[idx], 'List Number')

    @ classmethod
    def write_all_definitions(cls,
                              report_document: Document,
                              text_input_document: Document,
                              text_input_soup: BeautifulSoup,
                              definitions_document: Document,
                              list_of_tables: List[str]
                              ):
        """
        Write the whole chapter 'Terms definitions'.

        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            definitions_document: .docx file where all definitions are written.
            list_of_tables: List of all table names.
        """

        definitions = cls(report_document, text_input_document, text_input_soup, definitions_document, list_of_tables)

        report_document.add_paragraph(definitions.TITLE, definitions.TITLE_STYLE)
        definitions.write_definitions()
