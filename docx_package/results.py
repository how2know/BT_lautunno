from docx.document import Document
from docx.shared import Cm
from typing import List, Dict, Union
from bs4 import BeautifulSoup

from docx_package import text_reading
from docx_package.picture import Picture


class ResultsChapter:
    """
    Class that represents the sub-chapter 'Discussion' of the chapters of the 'Results' section.
    """

    # name of the discussion sub-chapter
    DISCUSSION_TITLE = 'Discussion'

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 text_input_soup: BeautifulSoup,
                 title: str,
                 list_of_tables: List[str],
                 picture_paths_list: List[str],
                 parameters_dictionary: Dict[str, Union[str, int]]
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            title: Title of the chapter.
            list_of_tables: List of all table names.
            picture_paths_list: List of the path of all input pictures.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value).
        """

        self.report = report_document
        self.text_input = text_input_document
        self.text_input_soup = text_input_soup
        self.title = title
        self.tables = list_of_tables
        self.picture_paths = picture_paths_list
        self.parameters_dictionary = parameters_dictionary

    def chapter_heading_index(self) -> int:
        """
        Returns:
            Paragraph index of the chapter heading in the text input document.
        """

        for paragraph_index, paragraph in enumerate(self.text_input.paragraphs):
            if paragraph.text == self.title and 'Heading' in paragraph.style.name:
                return paragraph_index

    def analysis_heading_index(self) -> int:
        """
        Returns:
            Paragraph index of the sub-chapter 'Discussion' heading in the text input document.
        """

        previous_index = self.chapter_heading_index()
        for paragraph_index, paragraph in enumerate(self.text_input.paragraphs[previous_index + 1:]):
            if paragraph.text == self.DISCUSSION_TITLE and 'Heading' in paragraph.style.name:
                return paragraph_index + previous_index + 1

    def next_heading_index(self) -> int:
        """
        Returns:
            Paragraph index of the following heading in the text input document.
        """

        previous_index = self.analysis_heading_index()
        for paragraph_index, paragraph in enumerate(self.text_input.paragraphs[previous_index + 1:]):
            if 'Heading' in paragraph.style.name:
                return paragraph_index + previous_index + 1

    @ property
    def paragraphs(self) -> List[str]:
        """
        Returns:
            List of all paragraphs (as text string) of the sub-chapter 'Discussion'.
        """

        list_of_paragraphs = []
        heading_index = self.analysis_heading_index()
        next_heading_index = self.next_heading_index()

        # store all paragraphs that are between the two heading indexes in a list
        for paragraph in self.text_input.paragraphs[heading_index + 1: next_heading_index]:
            list_of_paragraphs.append(paragraph.text)

        return list_of_paragraphs

    @ property
    def chapter_parameters(self) -> List[str]:
        """
        Returns:
            List of parameters needed to be writen in the sub-chapter.
        """

        return text_reading.get_dropdown_list_of_table(self.text_input_soup,
                                                       self.tables.index('{} parameter table'.format(self.title))
                                                       )

    @property
    def picture_name(self) -> str:
        """
        Returns:
            Title with underscores instead of spaces, e.g. 'Use environment' becomes 'Use_environment'.
        """

        return self.title.replace(' ', '_')

    @property
    def picture_captions(self) -> List[str]:
        """
        Returns:
            List of the captions of the pictures.
        """

        # read the caption text from the corresponding table in text input and append it to a list
        captions_list = []
        table_index = self.tables.index('{} caption table'.format(self.title))
        table = self.text_input.tables[table_index]
        for i in range(1, 4):
            cell = table.cell(i, 1)
            captions_list.append(cell.text)

        return captions_list

    def add_picture(self):
        """
        Add max. 3 pictures to the chapter.
        """

        captions = self.picture_captions
        picture_name = self.picture_name

        # add pictures that correspond to the given picture file names with the corresponding captions
        for i in range(0, 3):
            Picture.add_picture_and_caption(self.report,
                                            self.picture_paths,
                                            picture_name + str(i + 1),
                                            captions[i],
                                            width=Cm(10)
                                            )

    def write_chapter(self):
        """
        Write the paragraphs of the sub-chapter including the parameters, the pictures and their caption.
        """

        # stores values of corresponding parameter keys in a list
        parameters_values = ['', '', '']
        for parameter_idx, parameter in enumerate(self.chapter_parameters):
            if parameter != '-':
                parameters_values[parameter_idx] = self.parameters_dictionary[parameter]

        # write paragraphs including values of parameters
        for paragraph in self.paragraphs:
            new_paragraph = self.report.add_paragraph(
                paragraph.format(parameters_values[0], parameters_values[1], parameters_values[2],)
            )
            new_paragraph.style.name = 'Normal'

        # add pictures and their caption at the end of the sub-chapter
        self.add_picture()
