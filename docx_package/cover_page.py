from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm
from typing import List, Dict, Union
import numpy as np
from PIL import Image, UnidentifiedImageError

from docx_package.layout import Layout
from docx_package.picture import Picture


class CoverPage:
    """
    Class that represents and creates the cover page of the report.
    """

    # cover page table for the picture caption
    COVER_PAGE_TABLE = 'Cover page table'

    # title and subtitle parameter key
    TITLE_KEY = 'Title'
    SUBTITLE_KEY = 'Subtitle'

    # keys for the approval table
    AUTHOR_NAME_KEY = 'Author’s name'
    AUTHOR_FUNCTION_KEY = 'Author’s function'
    REVIEWER_NAME_KEY = 'Reviewer’s name'
    REVIEWER_FUNCTION_KEY = 'Reviewer’s function'
    APPROVER_NAME_KEY = 'Approver’s name'
    APPROVER_FUNCTION_KEY = 'Approver’s function'

    # hexadecimal of color for cell shading
    LIGHT_GREY_10 = 'D0CECE'

    # picture file name without the extension
    PICTURE_NAME = 'Cover_page'

    # width of the picture of the cover page
    PICTURE_WIDTH = Cm(14)

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 list_of_tables: List[str],
                 picture_paths_list: List[str],
                 parameters_dictionary: Dict[str, Union[str, int]]):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            list_of_tables: List of all table names.
            picture_paths_list: List of the path of all input pictures.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value).
        """

        self.report = report_document
        self.text_input = text_input_document
        self.tables = list_of_tables
        self.picture_paths = picture_paths_list
        self.parameters = parameters_dictionary

    def write_title(self):
        """
        Create and add a title with a first capital letter and border below it.
        """

        title_text = Layout.capitalize_first_letter(self.parameters[self.TITLE_KEY])
        title = self.report.add_paragraph(title_text, 'Title')
        Layout.insert_horizontal_border(title)

    def write_subtitle(self) -> Paragraph:
        """
        Create and add a subtitle with a first capital letter.

        Returns:
            Paragraph of the subtitle.
        """

        subtitle_text = Layout.capitalize_first_letter(self.parameters[self.SUBTITLE_KEY])
        subtitle = self.report.add_paragraph(subtitle_text, 'Subtitle')
        return subtitle

    # TODO: set column size
    def add_approval_table(self):
        """
        Add a table for the approval of the document.
        """

        # name of the different persons
        author_name = self.parameters[self.AUTHOR_NAME_KEY]
        reviewer_name = self.parameters[self.REVIEWER_NAME_KEY]
        approver_name = self.parameters[self.APPROVER_NAME_KEY]

        # store content of cells in a matrix
        approval_cells_text = np.array((['Role', 'Name / Function', 'Date', 'Signature'],
                                        ['Author', author_name, '', ''],
                                        ['Reviewer', reviewer_name, '', ''],
                                        ['Approver', approver_name, '', '']))

        # create table, define its style and fill it
        approval_table = self.report.add_table(rows=4, cols=4)
        # layout.define_table_style(approval_table)
        approval_table.style = 'Table Grid'
        # approval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        approval_table.autofit = True
        for i in range(0, 4):
            for j in range(0, 4):
                approval_table.cell(i, j).text = approval_cells_text[i, j]

        # set the shading of the first row to light_grey_10 and make it bold
        for cell in approval_table.rows[0].cells:
            Layout.set_cell_shading(cell, self.LIGHT_GREY_10)
            cell.paragraphs[0].runs[0].font.bold = True

        # function of the different persons
        author_function = Layout.capitalize_first_letter(self.parameters[self.AUTHOR_FUNCTION_KEY])
        reviewer_function = Layout.capitalize_first_letter(self.parameters[self.REVIEWER_FUNCTION_KEY].capitalize())
        approver_function = Layout.capitalize_first_letter(self.parameters[self.APPROVER_FUNCTION_KEY].capitalize())

        # add the function to the person
        approval_table.cell(1, 1).add_paragraph(author_function)
        approval_table.cell(2, 1).add_paragraph(reviewer_function)
        approval_table.cell(3, 1).add_paragraph(approver_function)

        # make the function italic or do nothing when no function were given
        try:
            for i in range(1, 4):
                approval_table.rows[i].cells[1].paragraphs[1].runs[0].font.italic = True
        except IndexError:
            pass

        # set the vertical alignment of all cells
        for row in approval_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    @ property
    def picture_caption(self) -> str:
        """
        Returns:
            The text of the caption of the cover page picture.
        """

        # find the text in the corresponding table in the text input document
        table_index = self.tables.index(self.COVER_PAGE_TABLE)
        table = self.text_input.tables[table_index]
        caption = table.cell(1, 1).text

        return caption

    def add_picture(self) -> bool:
        """
        Load a picture from the input files and add it to the report.

        The longest side (height or width) is set to 14 cm and the ratio is kept.
        The picture is added in the center regarding the side margin and spacing at the top and the bottom
        of the picture is set according to the height.

        Returns:
            True if a picture was added, and False if not.
        """

        # find the files that correspond to the picture file name
        for picture_path in self.picture_paths:
            if self.PICTURE_NAME in picture_path:

                try:
                    picture = Image.open(picture_path)

                    # find the longest side and set it to 14 cm when adding the picture
                    # case where the width is the longest side
                    if picture.width >= picture.height:

                        # set the spacing before and after the picture according the height/width ratio
                        if picture.height / picture.width * 14 < 5:
                            space = Cm(5)
                        elif picture.height / picture.width * 14 < 10:
                            space = Cm(3)
                        elif picture.height / picture.width * 14 < 14:
                            space = Cm(1)

                        # add the picture and its caption
                        Picture.add_picture_and_caption(self.report,
                                                        self.picture_paths,
                                                        self.PICTURE_NAME,
                                                        self.picture_caption,
                                                        width=self.PICTURE_WIDTH,
                                                        space_before=space,
                                                        space_after=space
                                                        )

                    # case where the height is the longest side
                    else:
                        # spacing before and after the picture is always 1 cm because height is always 14 cm
                        space = Cm(1)

                        # add the picture and its caption
                        Picture.add_picture_and_caption(self.report,
                                                        self.picture_paths,
                                                        self.PICTURE_NAME,
                                                        self.picture_caption,
                                                        width=self.PICTURE_WIDTH,
                                                        space_before=space,
                                                        space_after=space
                                                        )

                    # terminate because a picture was added and return True
                    return True

                # do nothing if the file is not an image
                except UnidentifiedImageError:
                    pass

        # return False because no picture was added
        return False

    def create(self):
        """
        Create the cover page with a title, a subtitle, a picture and its caption
        and a table for the approval of the report.
        """

        # add title, subtitle, picture (if an image file is provided) and approval table
        self.write_title()
        subtitle = self.write_subtitle()
        picture_added = self.add_picture()
        self.add_approval_table()

        # set the spacing between subtitle and approval table to 16 cm if no picture was added
        if not picture_added:
            subtitle.paragraph_format.space_after = Cm(16)






